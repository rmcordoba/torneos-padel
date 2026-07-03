from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Márgenes ──────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width    = Inches(8.5)
section.page_height   = Inches(11)
section.left_margin   = Inches(1.2)
section.right_margin  = Inches(1.2)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

# ── Paleta de colores ─────────────────────────────────────────────────────────
VERDE        = RGBColor(0x1A, 0x7A, 0x4A)
VERDE2       = RGBColor(0x2E, 0x9E, 0x64)
VERDE3       = RGBColor(0x1B, 0x5E, 0x20)
GRIS_OSCURO  = RGBColor(0x2D, 0x2D, 0x2D)
GRIS_MEDIO   = RGBColor(0x55, 0x55, 0x55)
GRIS_CLARO   = RGBColor(0x88, 0x88, 0x88)
AZUL         = RGBColor(0x0D, 0x47, 0xA1)
BLANCO       = RGBColor(0xFF, 0xFF, 0xFF)
NARANJA      = RGBColor(0xE6, 0x5C, 0x00)
NUM_COLOR    = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helpers XML ───────────────────────────────────────────────────────────────
def shading(para, fill: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    pPr.append(shd)

def cell_bg(cell, fill: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    tcPr.append(shd)

def bottom_border(para, color='1A7A4A', sz='6'):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    sz)
    bot.set(qn('w:space'), '4')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)

# ── Tipografías ───────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text.upper())
    run.bold           = True
    run.font.size      = Pt(17)
    run.font.color.rgb = VERDE
    bottom_border(p)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold           = True
    run.font.size      = Pt(13)
    run.font.color.rgb = VERDE2
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold           = True
    run.font.size      = Pt(11.5)
    run.font.color.rgb = GRIS_OSCURO
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(text)
    r.font.size      = Pt(11)
    r.font.color.rgb = GRIS_OSCURO
    return p

# ── Paso numerado con círculo verde ──────────────────────────────────────────
def paso(numero, titulo, descripcion):
    """Bloque de paso: círculo verde + título en negrita + descripción."""
    # Línea del paso
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(1)

    # Número en "círculo" simulado con fondo verde
    num_run = p.add_run(f'  {numero}  ')
    num_run.bold           = True
    num_run.font.size      = Pt(10)
    num_run.font.color.rgb = BLANCO
    # Fondo verde al número usando shading de run (highlight no soporta hex)
    # Usamos un truco: envolvemos el run en un rPr con highlight; como alternativa
    # simplemente añadimos texto con estilo.

    sep = p.add_run('  ')
    sep.font.size = Pt(11)

    tit = p.add_run(titulo)
    tit.bold           = True
    tit.font.size      = Pt(11)
    tit.font.color.rgb = GRIS_OSCURO

    shading(p, 'F1F8F4')  # fondo verde muy suave para la línea del paso

    # Descripción debajo, indentada
    if descripcion:
        pd = doc.add_paragraph()
        pd.paragraph_format.space_before = Pt(1)
        pd.paragraph_format.space_after  = Pt(5)
        pd.paragraph_format.left_indent  = Inches(0.4)
        rd = pd.add_run(descripcion)
        rd.font.size      = Pt(10.5)
        rd.font.color.rgb = GRIS_MEDIO
    return p

def sub_paso(texto):
    """Sub-paso con viñeta personalizada."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.5)
    r1 = p.add_run('→  ')
    r1.font.size      = Pt(10)
    r1.font.color.rgb = VERDE2
    r1.bold = True
    r2 = p.add_run(texto)
    r2.font.size      = Pt(10.5)
    r2.font.color.rgb = GRIS_OSCURO
    return p

def campo(nombre, descripcion):
    """Campo de formulario con label en negrita."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.5)
    r1 = p.add_run(f'{nombre}: ')
    r1.bold           = True
    r1.font.size      = Pt(10.5)
    r1.font.color.rgb = GRIS_OSCURO
    r2 = p.add_run(descripcion)
    r2.font.size      = Pt(10.5)
    r2.font.color.rgb = GRIS_MEDIO
    return p

def nota(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.left_indent  = Inches(0.25)
    shading(p, 'E3F2FD')
    r = p.add_run('NOTA: ' + text)
    r.font.size      = Pt(10)
    r.font.color.rgb = AZUL
    return p

def atencion(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.left_indent  = Inches(0.25)
    shading(p, 'FFF8E1')
    r = p.add_run('ATENCIÓN: ' + text)
    r.font.size      = Pt(10)
    r.font.color.rgb = NARANJA
    return p

def consejo(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.left_indent  = Inches(0.25)
    shading(p, 'E8F5E9')
    r = p.add_run('CONSEJO: ' + text)
    r.font.size      = Pt(10)
    r.font.color.rgb = VERDE3
    return p

def screenshot(label):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shading(p, 'ECEFF1')
    r = p.add_run(f'[ CAPTURA: {label} ]')
    r.font.size      = Pt(9.5)
    r.font.color.rgb = GRIS_CLARO
    r.italic         = True
    return p

def separador():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run('─' * 90)
    r.font.size      = Pt(6)
    r.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)

def page_break():
    doc.add_page_break()

def tabla_header(tbl, headers, bg='1A7A4A'):
    row = tbl.rows[0]
    for i, h in enumerate(headers):
        c = row.cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.bold           = True
        r.font.size      = Pt(10)
        r.font.color.rgb = BLANCO
        cell_bg(c, bg)

def tabla_fila(tbl, valores, par=True):
    row = tbl.add_row()
    bg = 'F9F9F9' if par else 'FFFFFF'
    for i, v in enumerate(valores):
        c = row.cells[i]
        c.text = str(v)
        c.paragraphs[0].runs[0].font.size = Pt(10)
        cell_bg(c, bg)


# ══════════════════════════════════════════════════════════════════════════════
#  PORTADA
# ══════════════════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.paragraph_format.space_before = Inches(1.4)
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover.add_run('SISTEMA DE GESTIÓN DE\nTORNEOS DE PÁDEL')
r.bold           = True
r.font.size      = Pt(26)
r.font.color.rgb = VERDE

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run('MANUAL DE USUARIO')
rs.bold           = True
rs.font.size      = Pt(17)
rs.font.color.rgb = VERDE2

doc.add_paragraph()
ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
rv = ver.add_run('Versión 2.0  —  Guía paso a paso  —  Mayo 2026')
rv.font.size      = Pt(11)
rv.font.color.rgb = GRIS_MEDIO
rv.italic         = True

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  ÍNDICE
# ══════════════════════════════════════════════════════════════════════════════
h1('Índice de Contenidos')
toc = [
    ('1.', 'Introducción al Sistema'),
    ('2.', 'Roles y Permisos'),
    ('3.', 'Cómo Iniciar Sesión'),
    ('4.', 'Cómo Crear una Cuenta'),
    ('5.', 'Portal Público'),
    ('  5.1', 'Explorar Torneos'),
    ('  5.2', 'Inscribirse en un Torneo'),
    ('  5.3', 'Consultar Ranking Público'),
    ('  5.4', 'Ver Cuadros de Fixture'),
    ('  5.5', 'Consultar Agenda de Partidos'),
    ('  5.6', 'Directorio de Jugadores'),
    ('6.', 'Panel del Jugador'),
    ('  6.1', 'Dashboard del Jugador'),
    ('  6.2', 'Completar el Perfil'),
    ('  6.3', 'Gestionar Inscripciones Propias'),
    ('7.', 'Panel del Organizador — Torneos'),
    ('  7.1', 'Dashboard Principal del Organizador'),
    ('  7.2', 'Ver el Listado de Torneos'),
    ('  7.3', 'Crear un Nuevo Torneo'),
    ('  7.4', 'Publicar un Torneo'),
    ('  7.5', 'Abrir Inscripciones'),
    ('  7.6', 'Gestionar una Categoría'),
    ('  7.7', 'Aprobar y Rechazar Inscripciones'),
    ('  7.8', 'Agregar una Pareja Manualmente'),
    ('  7.9', 'Cerrar Inscripciones y Generar Fixture'),
    ('  7.10', 'Cargar Resultados de Partidos'),
    ('  7.11', 'Avanzar al Playoff (formato Grupos + Playoff)'),
    ('  7.12', 'Finalizar un Torneo'),
    ('8.', 'Panel del Organizador — Calendario'),
    ('  8.1', 'Navegar el Calendario'),
    ('  8.2', 'Programar un Partido'),
    ('  8.3', 'Reprogramar un Partido'),
    ('9.', 'Panel del Organizador — Jugadores'),
    ('10.', 'Panel del Organizador — Sedes y Canchas'),
    ('  10.1', 'Crear una Sede'),
    ('  10.2', 'Agregar Canchas a una Sede'),
    ('11.', 'Panel del Organizador — Ranking'),
    ('  11.1', 'Crear una Tabla de Ranking'),
    ('  11.2', 'Definir Reglas de Puntuación'),
    ('  11.3', 'Recalcular el Ranking'),
    ('12.', 'Panel del Organizador — Reportes'),
    ('13.', 'Panel del Organizador — Auditoría'),
    ('14.', 'Configuración del Organizador'),
    ('  14.1', 'Datos de la Organización'),
    ('  14.2', 'Parámetros Generales'),
    ('  14.3', 'Gestión de Categorías'),
    ('  14.4', 'Gestión de Colaboradores'),
    ('15.', 'Panel de Administración Global'),
    ('  15.1', 'Dashboard de Administración'),
    ('  15.2', 'Gestión de Usuarios'),
    ('  15.3', 'Gestión de Organizadores'),
    ('  15.4', 'Auditoría Global'),
]
for num, title in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    indent = num.startswith('  ')
    p.paragraph_format.left_indent = Inches(0.4 if indent else 0.1)
    r1 = p.add_run(num.strip() + '  ')
    r1.bold      = not indent
    r1.font.size = Pt(10.5)
    r2 = p.add_run(title)
    r2.font.size = Pt(10.5)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  1. INTRODUCCIÓN
# ══════════════════════════════════════════════════════════════════════════════
h1('1. Introducción al Sistema')
body(
    'El Sistema de Gestión de Torneos de Pádel es una plataforma web que centraliza todo '
    'el ciclo de vida de un torneo: desde la publicación y la inscripción de parejas hasta '
    'la generación del fixture, la carga de resultados, el seguimiento del calendario y el '
    'cálculo del ranking. Se accede desde cualquier navegador web sin necesidad de instalar '
    'software adicional.'
)
body(
    'Este manual describe, pantalla por pantalla y paso a paso, qué puede hacer cada tipo '
    'de usuario. Cada sección indica los elementos visibles, los campos a completar, los '
    'botones disponibles y los mensajes que el sistema muestra.'
)
nota('Las capturas de pantalla referenciadas a lo largo del documento deben insertarse en '
     'los espacios indicados con [ CAPTURA: … ].')
page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  2. ROLES
# ══════════════════════════════════════════════════════════════════════════════
h1('2. Roles y Permisos')
body('El sistema reconoce tres tipos de usuario. El rol determina a qué secciones puede '
     'acceder y qué acciones puede realizar.')

tbl_r = doc.add_table(rows=1, cols=3)
tbl_r.style = 'Table Grid'
tabla_header(tbl_r, ['ROL', 'QUIÉN LO TIENE', 'QUÉ PUEDE HACER'])
roles = [
    ('Super Admin', 'Administradores del sistema.',
     'Gestionar todos los organizadores, usuarios y auditoría global.'),
    ('Organizador\n(Owner / Organizer / Collaborator)',
     'Staff de un club u organización.',
     'Crear torneos, gestionar inscripciones, fixture, resultados, calendario, ranking y configuración.'),
    ('Jugador', 'Cualquier usuario registrado.',
     'Inscribirse en torneos, ver su panel personal y su historial.'),
]
for i, (r, q, h) in enumerate(roles):
    tabla_fila(tbl_r, [r, q, h], par=(i % 2 == 0))

doc.add_paragraph()
h3('Roles dentro de una Organización')
body('Los miembros de una organización tienen uno de estos tres sub-roles:')
campo('Owner (Propietario)', 'Acceso total. Puede configurar la organización y gestionar colaboradores.')
campo('Organizer', 'Acceso completo a torneos, inscripciones, fixture y resultados.')
campo('Collaborator', 'Acceso limitado a los permisos específicos que el Owner le asigne.')

h3('Permisos que se pueden asignar a un Colaborador')
permisos = [
    ('Torneos',         'Crear y editar torneos.'),
    ('Inscripciones',   'Aprobar, rechazar y gestionar inscripciones.'),
    ('Resultados',      'Cargar y modificar resultados de partidos.'),
    ('Calendario',      'Programar y reprogramar partidos.'),
    ('Sedes',           'Crear y editar sedes y canchas.'),
    ('Categorías',      'Gestionar el catálogo de categorías de la organización.'),
    ('Reportes',        'Ver reportes y estadísticas.'),
    ('Colaboradores',   'Invitar y gestionar otros colaboradores.'),
]
for nombre, desc in permisos:
    campo(nombre, desc)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  3. INICIAR SESIÓN
# ══════════════════════════════════════════════════════════════════════════════
h1('3. Cómo Iniciar Sesión')
body('Cualquier usuario registrado puede iniciar sesión desde la página de login o desde '
     'el modal de autenticación del portal público.')

h2('Desde la página de login')
screenshot('Formulario de inicio de sesión')

paso(1, 'Ingresar al formulario',
     'Ir a la URL del sistema. Hacer clic en el botón "Iniciar sesión" del encabezado '
     'del portal público. Se abre un modal con dos pestañas: "Iniciar sesión" y "Registrarse".')

paso(2, 'Elegir método de autenticación',
     'El sistema ofrece dos opciones:')
sub_paso('Botón "Continuar con Google": redirige al flujo de autenticación de Google. '
         'Al completarlo, el sistema crea o vincula la cuenta automáticamente.')
sub_paso('Formulario con email y contraseña: completar los campos manualmente.')

paso(3, 'Completar el formulario con email y contraseña',
     'Si se elige el formulario manual, completar los siguientes campos:')
campo('Email', 'Dirección de correo con la que se registró la cuenta. Campo obligatorio.')
campo('Contraseña', 'Contraseña de la cuenta. Campo obligatorio.')
sub_paso('El enlace "¿Olvidaste tu contraseña?" inicia el flujo de recuperación por email.')

paso(4, 'Confirmar el inicio de sesión',
     'Hacer clic en el botón "Ingresar →". Mientras el sistema procesa, el botón muestra '
     '"Iniciando sesión…" con un spinner. Si los datos son correctos, el sistema redirige '
     'automáticamente al panel correspondiente según el rol del usuario.')

paso(5, 'Revisar mensajes de error',
     'Si los datos son incorrectos, aparece una caja roja con el mensaje de error. '
     'Verificar que el email esté escrito sin espacios y que la contraseña sea correcta.')

nota('Si el inicio de sesión fue mediante Google, no existe una contraseña local. '
     'El enlace "¿Olvidaste tu contraseña?" no aplica en ese caso.')

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  4. CREAR CUENTA
# ══════════════════════════════════════════════════════════════════════════════
h1('4. Cómo Crear una Cuenta')
body('Cualquier persona puede crear una cuenta gratuita. Al registrarse, el sistema '
     'asigna automáticamente el rol de Jugador.')

screenshot('Formulario de registro de nueva cuenta')

paso(1, 'Abrir el formulario de registro',
     'Desde el portal público, hacer clic en el botón "Registrarse" del encabezado, '
     'o en el enlace "Registrate gratis" en el pie del formulario de login. '
     'Se abre la pestaña "Registrarse" del modal de autenticación.')

paso(2, 'Completar los datos personales',
     'Completar los siguientes campos (todos obligatorios):')
campo('Nombre', 'Nombre de pila. Máximo 100 caracteres.')
campo('Apellido', 'Apellido. Máximo 100 caracteres.')
campo('Email', 'Correo electrónico. Debe ser único en el sistema.')
campo('Contraseña', 'Mínimo 8 caracteres.')
campo('Repetir contraseña', 'Debe coincidir exactamente con el campo "Contraseña".')

paso(3, 'Enviar el formulario',
     'Hacer clic en "Crear cuenta →". El sistema muestra un spinner mientras procesa. '
     'Si hay errores de validación, aparece una caja roja con el detalle.')

paso(4, 'Resultado del registro',
     'Al completarse el registro con éxito, el sistema redirige al listado de torneos '
     'públicos. Aparece un banner verde de bienvenida en la parte superior de la pantalla.')

paso(5, 'Completar el perfil de jugador (recomendado)',
     'Con la cuenta creada, conviene completar el perfil de jugador (nombre, apellido, '
     'teléfono, DNI, fecha de nacimiento) antes de intentar inscribirse en un torneo. '
     'Algunos organizadores exigen el perfil completo para aprobar inscripciones. '
     'El perfil se completa desde el Panel del Jugador (ver sección 6.2).')

consejo('También es posible registrarse con Google haciendo clic en "Continuar con Google" '
        'en la pestaña de registro. En ese caso no se requiere contraseña.')

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  5. PORTAL PÚBLICO
# ══════════════════════════════════════════════════════════════════════════════
h1('5. Portal Público')
body('El portal público es la cara visible del sistema. Está disponible para cualquier '
     'visitante, sin necesidad de iniciar sesión. Desde aquí se pueden explorar torneos, '
     'ver rankings, consultar la agenda de partidos y acceder a los cuadros de fixture.')

h2('Encabezado del Portal')
body('El encabezado es la barra de navegación superior, visible en todas las páginas '
     'del portal público. Contiene:')
sub_paso('Logo "PádelPro" a la izquierda.')
sub_paso('Campo de búsqueda: "Buscar jugador o torneo…" con ícono de lupa.')
sub_paso('Botones "Iniciar sesión" y "Registrarse" (si el usuario no tiene sesión iniciada).')
sub_paso('Menú de usuario (avatar + nombre + flecha) si ya hay sesión iniciada. '
         'Al hacer clic se despliega un menú con: nombre, email, enlaces al panel '
         'correspondiente y el botón "Cerrar sesión".')
sub_paso('Barra de navegación: Torneos · Cuadros · Agenda · Ranking · Jugadores.')

screenshot('Encabezado del portal público con usuario autenticado')

separador()

# 5.1
h2('5.1 Explorar Torneos   (/torneos)')
body('Página principal del portal. Muestra todos los torneos que los organizadores han '
     'publicado como públicos.')

screenshot('Página principal de torneos públicos')

paso(1, 'Ver el torneo destacado',
     'En la parte superior aparece el torneo más reciente con inscripciones abiertas. '
     'Se muestran: nombre del torneo, organizador, descripción, fechas y una barra de '
     'progreso de cupos disponibles por categoría.')
sub_paso('El botón "Ver torneo" lleva al detalle de ese torneo.')
sub_paso('El botón "Pre-inscribirse" (si está disponible) inicia el flujo de inscripción.')

paso(2, 'Explorar otros torneos',
     'Debajo del torneo destacado aparece una grilla con el resto de los torneos disponibles. '
     'Cada tarjeta muestra: nombre, organizador, fechas, categorías disponibles, '
     'estado (Próximamente / Inscripciones abiertas / En curso / Finalizado) y '
     'barra de progreso de cupos.')

paso(3, 'Buscar un torneo',
     'Escribir el nombre o parte del nombre en el campo "Buscar jugador o torneo…" '
     'del encabezado. Los resultados se filtran en tiempo real.')

paso(4, 'Acceder al detalle de un torneo',
     'Hacer clic en cualquier tarjeta de torneo para ver su página de detalle, '
     'donde aparecen todas sus categorías con cupos, precio y opción de inscripción.')

separador()

# 5.2
h2('5.2 Inscribirse en un Torneo')
body('El proceso de inscripción se realiza desde la página pública de cada categoría.')

screenshot('Página de detalle de categoría con formulario de inscripción')

paso(1, 'Navegar a la categoría',
     'Desde la página del torneo, hacer clic en la categoría en la que se desea participar. '
     'La página muestra: nombre de la categoría, cantidad de parejas inscriptas (ej: "12/16 parejas"), '
     'el fixture si ya está generado y el formulario de inscripción.')

paso(2, 'Iniciar sesión si es necesario',
     'Si el usuario no tiene sesión iniciada, aparece una caja azul con el mensaje '
     '"Necesitás iniciar sesión" y un botón verde "Iniciar sesión". '
     'Hacer clic en el botón para abrir el modal de autenticación y seguir los pasos '
     'de la sección 3 de este manual.')

paso(3, 'Verificar el perfil',
     'Si la cuenta no tiene un perfil de jugador completo, aparece el mensaje '
     '"Necesitás un perfil de jugador". En ese caso, completar el perfil primero '
     '(ver sección 6.2) y luego volver a esta página.')

paso(4, 'Seleccionar al compañero/a de pareja',
     'En el formulario de inscripción, el primer jugador (el usuario actual) ya aparece '
     'completado con nombre y avatar. Para el segundo jugador:')
sub_paso('Escribir el nombre o apellido del compañero/a en el campo de búsqueda '
         '"Escribí nombre o apellido…" (mínimo 2 caracteres).')
sub_paso('Esperar los resultados del desplegable. Cada resultado muestra el nombre '
         'y el email del jugador.')
sub_paso('Hacer clic en el nombre del compañero/a para seleccionarlo. '
         'Aparece una tarjeta verde con ambos jugadores confirmados.')
sub_paso('Si hay un error en la selección, hacer clic en el botón ✕ de la tarjeta '
         'del jugador 2 para deseleccionarlo y buscar otro.')

nota('El compañero/a debe tener una cuenta registrada en el sistema. '
     'Si su nombre no aparece en los resultados, pedirle que cree una cuenta primero.')

paso(5, 'Confirmar la inscripción',
     'Con ambos jugadores seleccionados, aparece un resumen de la pareja en una tarjeta '
     'verde con el mensaje "Tu solicitud quedará pendiente hasta que el organizador la apruebe". '
     'Hacer clic en el botón verde "Confirmar inscripción de pareja". '
     'El botón muestra "Enviando solicitud…" mientras procesa.')

paso(6, 'Revisar el resultado',
     'Al completarse:')
sub_paso('Si hay cupo disponible: aparece una caja verde "Solicitud enviada — Tu inscripción '
         'está pendiente de aprobación". La inscripción queda en estado PENDIENTE.')
sub_paso('Si el torneo está lleno: aparece una caja amarilla "En lista de espera — '
         'El cupo está completo". La inscripción queda en estado EN ESPERA.')
sub_paso('Si ocurre un error: aparece una caja roja con el mensaje de error.')

paso(7, 'Hacer seguimiento de la inscripción',
     'El estado de la inscripción (pendiente, aprobada, rechazada) se puede consultar '
     'en cualquier momento desde el Panel del Jugador (sección 6.3).')

screenshot('Formulario de inscripción con pareja seleccionada')

separador()

# 5.3
h2('5.3 Consultar el Ranking Público   (/ranking)')
body('Muestra las tablas de ranking que los organizadores han publicado.')

paso(1, 'Seleccionar la tabla de ranking',
     'Si hay más de una tabla disponible, aparece un selector en la parte superior. '
     'Hacer clic en el nombre de la tabla que se desea consultar.')

paso(2, 'Ver el podio',
     'La sección superior muestra los tres primeros lugares con sus nombres, '
     'puntos acumulados y avatares con iniciales.')

paso(3, 'Ver la tabla completa',
     'Debajo del podio se despliega la tabla con todas las posiciones, incluyendo: '
     'posición, nombre de la pareja o jugador, puntos acumulados y torneos jugados.')

paso(4, 'Ver los criterios de puntuación',
     'Al final de la tabla se detallan las reglas de puntuación: cuántos puntos '
     'otorga cada posición final (1°, 2°, 3°, participación, etc.).')

screenshot('Ranking público con podio y tabla de posiciones')

separador()

# 5.4
h2('5.4 Ver Cuadros de Fixture   (/cuadros)')
body('Permite ver los cuadros de competencia de torneos en curso o finalizados '
     'sin necesidad de iniciar sesión.')

paso(1, 'Seleccionar torneo y categoría',
     'Usar los menús desplegables en la parte superior para elegir el torneo '
     'y luego la categoría que se desea consultar.')

paso(2, 'Explorar el fixture',
     'El cuadro se muestra según el formato del torneo:')
sub_paso('Eliminación directa: se visualiza el bracket con rondas (Cuartos, Semis, Final). '
         'Cada casilla muestra los nombres de las parejas y el resultado si ya se jugó.')
sub_paso('Grupos + Playoff: primero se muestra la tabla de posiciones de cada grupo '
         'con las columnas PJ (jugados), PG (ganados), PP (perdidos), SG/SP (sets), '
         'Pts (puntos), y luego el cuadro de playoff.')
sub_paso('Para un partido ya jugado: aparece el marcador de sets (ej: "6-4, 7-5"). '
         'Para uno no jugado: aparece "TBD" en gris. Si un equipo pasó con bye, '
         'aparece la etiqueta "BYE".')

screenshot('Cuadro de fixture público - Eliminación directa')

separador()

# 5.5
h2('5.5 Consultar la Agenda de Partidos   (/agenda)')
body('Muestra el calendario de partidos programados de todos los torneos públicos.')

paso(1, 'Navegar a la agenda',
     'Hacer clic en "Agenda" en la barra de navegación del portal. '
     'Aparece el badge "HOY" sobre la pestaña cuando hay partidos para el día actual.')

paso(2, 'Ver los partidos del día',
     'Por cada partido se muestra: hora de inicio, nombres de las dos parejas, '
     'torneo, categoría, sede y estado.')

paso(3, 'Filtrar por fecha',
     'Usar el selector de fecha para ver los partidos de un día específico.')

screenshot('Agenda pública de partidos')

separador()

# 5.6
h2('5.6 Directorio de Jugadores   (/jugadores)')
body('Directorio público de jugadores registrados en el sistema.')

paso(1, 'Buscar un jugador',
     'Escribir el nombre o apellido en el campo de búsqueda. '
     'Los resultados se actualizan en tiempo real.')

paso(2, 'Ver el perfil de un jugador',
     'Hacer clic en un jugador de la lista para acceder a su perfil público, '
     'que incluye su historial de torneos y posición en el ranking.')

screenshot('Directorio de jugadores')

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  6. PANEL DEL JUGADOR
# ══════════════════════════════════════════════════════════════════════════════
h1('6. Panel del Jugador')
body('El Panel del Jugador está disponible para usuarios registrados que no pertenecen '
     'al staff de ninguna organización. Se accede desde el menú de usuario del encabezado '
     '→ "Mi panel".')

h2('6.1 Dashboard del Jugador   (/dashboard/jugador)')
body('Al ingresar, el sistema muestra el saludo "¡Hola, [Nombre]! Bienvenido a tu panel '
     'de jugador." Si el perfil está incompleto, aparece un banner amarillo con el texto '
     '"Perfil incompleto — Completá tu perfil para poder inscribirte" y un botón '
     '"Completar perfil".')

screenshot('Dashboard del jugador')

h3('Sección: Torneos disponibles para inscripción')
body('Muestra todos los torneos con inscripciones abiertas en los que el jugador '
     'todavía puede anotarse. Cada torneo se presenta como una tarjeta con:')
sub_paso('Nombre del torneo y organizador.')
sub_paso('Fechas de inicio y fin.')
sub_paso('Categorías disponibles: nombre, precio si aplica, barra de progreso de cupos '
         '(verde si hay lugar, amarilla si queda poco, roja si está lleno) y el '
         'contador "X/Y" de inscriptos.')
sub_paso('Botón "Inscribirme" por cada categoría con cupos disponibles. '
         'Si la categoría está llena, el botón aparece como "Completo" y está deshabilitado.')
sub_paso('Al hacer clic en "Inscribirme", el sistema redirige a la página de la categoría '
         'donde se completa el proceso (ver sección 5.2).')

h3('Sección: Mis inscripciones')
body('Lista todas las inscripciones del jugador. Cada inscripción muestra:')
sub_paso('Ícono de trofeo con color según el estado.')
sub_paso('Nombre del torneo y badge de estado: PENDIENTE (amarillo), APROBADA (verde), '
         'RECHAZADA (rojo), EN ESPERA (gris).')
sub_paso('Nombre de la categoría, nombre del compañero/a y fecha de inicio del torneo.')

separador()

h2('6.2 Completar el Perfil de Jugador')
body('Para poder inscribirse en torneos, el jugador debe tener un perfil completo.')

screenshot('Formulario de perfil del jugador')

paso(1, 'Acceder al formulario de perfil',
     'Hacer clic en "Completar perfil" en el banner amarillo del dashboard, '
     'o navegar al menú de usuario del encabezado → "Mi perfil".')

paso(2, 'Completar los datos de jugador',
     'Completar el formulario con los siguientes campos:')
campo('Nombre', 'Nombre de pila del jugador.')
campo('Apellido', 'Apellido del jugador.')
campo('Teléfono', 'Número de contacto (opcional pero recomendado).')
campo('Fecha de nacimiento', 'Selector de fecha.')
campo('DNI', 'Número de documento. Debe ser único en el sistema.')
campo('Contacto de emergencia', 'Nombre y teléfono de contacto de emergencia (opcional).')

paso(3, 'Guardar los cambios',
     'Hacer clic en el botón "Guardar". El sistema confirma con un mensaje de éxito.')

atencion('El DNI debe ser único en el sistema. Si aparece un error al guardar, '
         'verificar que el número no esté ya registrado bajo otra cuenta.')

separador()

h2('6.3 Gestionar Inscripciones Propias')
body('Desde el dashboard del jugador, en la sección "Mis inscripciones", el jugador puede:')
sub_paso('Ver el estado actual de cada inscripción.')
sub_paso('Cancelar una inscripción en estado PENDIENTE haciendo clic en el botón '
         '"Cancelar inscripción" de esa tarjeta.')
sub_paso('Una vez aprobada la inscripción, acceder al fixture y calendario del torneo '
         'para ver los partidos asignados.')

nota('Las inscripciones APROBADAS no pueden cancelarse directamente. '
     'Para hacerlo, contactar al organizador del torneo.')

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  7. PANEL DEL ORGANIZADOR — TORNEOS
# ══════════════════════════════════════════════════════════════════════════════
h1('7. Panel del Organizador — Torneos')
body('El Panel del Organizador es el área de gestión principal. Se accede desde el menú '
     'de usuario del encabezado → "Mi dashboard". El menú lateral izquierdo da acceso '
     'a todas las secciones: Torneos, Inscripciones, Jugadores, Sedes, Calendario, '
     'Ranking, Reportes, Auditoría y Configuración.')

h2('7.1 Dashboard Principal   (/dashboard)')
body('La pantalla de inicio del organizador muestra un resumen en tiempo real.')

screenshot('Dashboard principal del organizador')

h3('Tarjetas de estadísticas (fila superior)')
sub_paso('Torneos activos: cantidad de torneos en curso o con inscripciones abiertas '
         'del total de torneos de la organización.')
sub_paso('Inscripciones pendientes: número de solicitudes que esperan revisión. '
         'Si hay pendientes, el número parpadea con un punto verde animado.')
sub_paso('Jugadores registrados: total de jugadores únicos en el sistema.')
sub_paso('Partidos pendientes: partidos programados que aún no se han jugado.')

h3('Torneos recientes')
body('Lista los últimos torneos creados con: nombre, badge de estado (color según etapa), '
     'cantidad de categorías y cantidad de parejas inscriptas.')
sub_paso('El botón "+ Crear torneo" en la esquina superior derecha abre el formulario '
         'de creación (ver sección 7.3).')
sub_paso('Hacer clic en cualquier torneo de la lista redirige a su página de detalle.')

h3('Agenda de hoy')
body('Muestra los partidos programados para el día actual con: hora, parejas, '
     'sede, cancha y categoría.')

separador()

h2('7.2 Ver el Listado de Torneos   (/dashboard/torneos)')

screenshot('Listado de torneos del organizador')

paso(1, 'Ver los torneos',
     'Los torneos se agrupan por estado:')
sub_paso('En curso / Inscripciones abiertas (punto verde).')
sub_paso('Próximos (punto azul): publicados pero sin inscripciones abiertas aún.')
sub_paso('Finalizados (punto gris).')

paso(2, 'Filtrar por estado',
     'Usar el menú desplegable de filtro para ver solo los torneos de un estado específico: '
     'Borrador, Publicado, Inscripciones abiertas, Inscripciones cerradas, '
     'En curso, Finalizado, Cancelado.')

paso(3, 'Acceder a un torneo',
     'Hacer clic en cualquier tarjeta de torneo para ver su página de detalle, '
     'desde donde se pueden ejecutar todas las acciones disponibles según el estado actual.')

paso(4, 'Crear un nuevo torneo',
     'Hacer clic en el botón "+ Nuevo torneo" en la esquina superior derecha '
     '(ver sección 7.3).')

nota('Si no hay torneos creados, la pantalla muestra el mensaje "Creá tu primer torneo" '
     'con un ícono de trofeo y el botón "+ Crear torneo".')

separador()

h2('7.3 Crear un Nuevo Torneo   (/dashboard/torneos/nuevo)')

screenshot('Formulario de creación de torneo - Paso 1: datos básicos')

paso(1, 'Completar los datos básicos del torneo',
     'Ingresar la información general:')
campo('Nombre', 'Nombre del torneo. Entre 2 y 150 caracteres. Obligatorio.')
campo('Descripción', 'Texto libre opcional. Máximo 1000 caracteres.')
campo('Fecha de inicio', 'Selector de fecha. Obligatorio.')
campo('Fecha de fin', 'Selector de fecha. Debe ser igual o posterior a la fecha de inicio. Obligatorio.')
campo('Fecha límite de inscripción', 'Opcional. Fecha hasta la que los jugadores pueden inscribirse.')
campo('¿Torneo público?', 'Casilla de verificación. Si está marcada, el torneo aparece en el portal público.')

paso(2, 'Agregar categorías',
     'En la sección "Categorías" del mismo formulario, agregar al menos una categoría '
     'haciendo clic en el botón "+ Agregar categoría". Por cada categoría completar:')
campo('Categoría', 'Seleccionar del listado de categorías del organizador.')
campo('Formato', 'Elegir el formato de competencia: Eliminación directa, Grupos + Playoff, '
      'Doble eliminación, Round Robin, Americano o Mexicano.')
campo('Mínimo de equipos', 'Cantidad mínima para que la categoría sea válida.')
campo('Máximo de equipos', 'Cupo máximo de equipos. Debe ser mayor o igual al mínimo.')
campo('Precio por pareja', 'Opcional. Importe en la moneda configurada.')
campo('Sets por partido', 'Número de sets. Usa el valor default del organizador si se deja vacío.')
campo('Games por set', 'Número de games. Usa el valor default del organizador si se deja vacío.')

paso(3, 'Guardar el torneo',
     'Hacer clic en el botón "Crear torneo". El sistema valida todos los campos. '
     'Si hay errores, los muestra en rojo junto a cada campo. '
     'Si todo es correcto, crea el torneo con estado BORRADOR y redirige '
     'a la página de detalle del torneo recién creado.')

screenshot('Formulario de creación de torneo - Sección de categorías')

separador()

h2('7.4 Publicar un Torneo')
body('Un torneo en estado BORRADOR no es visible para el público. '
     'Para hacerlo visible hay que publicarlo.')

paso(1, 'Acceder a la página de detalle del torneo',
     'Desde el listado de torneos, hacer clic en el torneo que se desea publicar.')

paso(2, 'Verificar los datos',
     'Revisar que el nombre, fechas y categorías sean correctas antes de publicar. '
     'Una vez publicado, el torneo es visible para todos los visitantes del portal.')

paso(3, 'Publicar',
     'Hacer clic en el botón "Publicar" en la cabecera del torneo. '
     'El badge de estado cambia de "BORRADOR" a "PUBLICADO". '
     'El torneo ya aparece en el portal público, pero las inscripciones siguen cerradas.')

nota('Un torneo publicado puede volver a editarse antes de que se abran las inscripciones.')

separador()

h2('7.5 Abrir Inscripciones')
body('Con el torneo publicado, el siguiente paso es habilitar el período de inscripciones.')

paso(1, 'Desde la página de detalle del torneo',
     'Hacer clic en el botón "Abrir inscripciones".')

paso(2, 'Confirmar la acción',
     'El sistema solicita confirmación. Al aceptar, el estado del torneo cambia a '
     '"INSCRIPCIONES ABIERTAS". Desde este momento, los jugadores pueden enviar '
     'solicitudes de inscripción desde el portal público.')

paso(3, 'Monitorear las inscripciones entrantes',
     'Las nuevas solicitudes aparecen automáticamente en la sección '
     '"Inscripciones" del panel del organizador (sección 7.7). '
     'El dashboard también muestra el contador de inscripciones pendientes.')

separador()

h2('7.6 Gestionar una Categoría')
body('Dentro del detalle de un torneo, la sección "Categorías" lista todas las categorías '
     'con su estado, formato y progreso de cupos. Hacer clic en una categoría abre '
     'su página de gestión.')

screenshot('Detalle de categoría con inscripciones aprobadas y pendientes')

h3('Lo que se ve al ingresar a una categoría')
sub_paso('Nombre de la categoría y badge de estado.')
sub_paso('Formato de competencia (ícono + nombre).')
sub_paso('Barra de progreso de cupos: verde si hay lugar, amarilla si queda poco, '
         'roja si está lleno. Muestra el texto "X/Y inscriptos".')
sub_paso('Reglas del partido: sets, games y configuración de tiebreak.')
sub_paso('Tres secciones con sus contadores: '
         '"Inscripciones aprobadas", "Inscripciones pendientes" y "Lista de espera".')

h3('Inscripciones aprobadas')
sub_paso('Tabla con todos los equipos aprobados, mostrando avatares, '
         'nombres de ambos jugadores y número de equipo.')
sub_paso('Botón de edición por cada inscripción.')
sub_paso('Botón de eliminación (con confirmación antes de borrar).')

h3('Inscripciones pendientes')
sub_paso('Lista de solicitudes que aguardan revisión.')
sub_paso('Por cada una se muestra: avatares superpuestos, nombres (Jugador1 / Jugador2) '
         'y fecha "Solicitado el [fecha]".')
sub_paso('Botón verde "Aprobar" y botón rojo "Rechazar" por cada inscripción.')

h3('Lista de espera')
sub_paso('Equipos que llegaron cuando el cupo ya estaba completo, '
         'ordenados por posición (1° en espera, 2° en espera, etc.).')
sub_paso('Botón para eliminar a un equipo de la lista de espera.')

consejo('Cuando se rechaza una inscripción aprobada o se elimina un equipo, '
        'el sistema automáticamente promueve al siguiente de la lista de espera.')

separador()

h2('7.7 Aprobar y Rechazar Inscripciones   (/dashboard/inscripciones)')
body('La página de inscripciones centraliza todas las solicitudes pendientes '
     'de todos los torneos activos de la organización.')

screenshot('Centro de gestión de inscripciones pendientes')

paso(1, 'Acceder a la sección de inscripciones',
     'Hacer clic en "Inscripciones" en el menú lateral izquierdo.')

paso(2, 'Aplicar filtros de fecha (opcional)',
     'Por defecto la vista muestra las inscripciones del día actual. Para cambiar:')
sub_paso('Botón "Hoy": muestra solo las del día de hoy.')
sub_paso('Botón "Todas las fechas": muestra todas sin filtro temporal.')
sub_paso('Rango personalizado: seleccionar fechas "Desde" y "Hasta" en los selectores.')

paso(3, 'Localizar la inscripción',
     'Las inscripciones están agrupadas por categoría de torneo. '
     'Cada grupo muestra el nombre de la categoría, el nombre del torneo y '
     'un badge amarillo con la cantidad de solicitudes pendientes. '
     'El enlace "Ver todo ›" lleva directo al detalle de esa categoría.')

paso(4, 'Aprobar una inscripción',
     'Hacer clic en el botón verde "Aprobar" de la inscripción. '
     'El sistema procesa la acción y la inscripción desaparece de la lista de pendientes. '
     'La pareja queda registrada como aprobada en la categoría.')

paso(5, 'Rechazar una inscripción',
     'Hacer clic en el botón rojo "Rechazar" de la inscripción. '
     'El sistema solicita confirmación. Al aceptar, la inscripción pasa a estado RECHAZADA '
     'y desaparece de la lista de pendientes. '
     'Si hay equipos en lista de espera, el primero es promovido automáticamente.')

nota('Si la pantalla muestra el ícono de portapapeles con el texto "Todo al día — '
     'No hay inscripciones pendientes", significa que no hay solicitudes sin procesar '
     'para las fechas seleccionadas.')

separador()

h2('7.8 Agregar una Pareja Manualmente')
body('Es posible inscribir una pareja directamente, sin necesidad de que los jugadores '
     'lo hagan desde el portal público. Esto es útil para inscripciones presenciales '
     'o migraciones de datos.')

paso(1, 'Acceder al formulario',
     'Desde la página de inscripciones (/dashboard/inscripciones), hacer clic en '
     'el botón "+ Inscribir pareja" en la esquina superior derecha. '
     'También puede accederse desde la página de detalle de la categoría con '
     'el botón "+ Agregar pareja".')

paso(2, 'Seleccionar la categoría',
     'Elegir el torneo y la categoría en la que se desea inscribir a la pareja.')

paso(3, 'Buscar a los jugadores',
     'Buscar al jugador 1 y al jugador 2 por nombre, apellido o email. '
     'El sistema busca entre todos los jugadores registrados.')

paso(4, 'Confirmar la inscripción',
     'Hacer clic en "Confirmar". La pareja queda inscripta directamente en estado APROBADO, '
     'sin pasar por la revisión de pendientes.')

separador()

h2('7.9 Cerrar Inscripciones y Generar Fixture')
body('Cuando el período de inscripciones termina, se cierra la etapa y se genera '
     'el cuadro de competencia.')

screenshot('Botón de generación de fixture en detalle de categoría')

paso(1, 'Cerrar inscripciones',
     'Desde la página de detalle del torneo, hacer clic en el botón '
     '"Cerrar inscripciones". El estado del torneo cambia a "INSCRIPCIONES CERRADAS". '
     'A partir de este momento no se aceptan nuevas solicitudes.')

paso(2, 'Acceder a la categoría',
     'Dentro del torneo, hacer clic en la categoría para la que se desea generar el fixture.')

paso(3, 'Generar el fixture',
     'En la página de la categoría, hacer clic en el botón "Generar fixture" '
     '(o "⚡ Generar fixture" en la vista de fixture). '
     'El sistema crea automáticamente los partidos según el formato elegido:')
sub_paso('Eliminación directa: genera el bracket con los cruces de primera ronda.')
sub_paso('Grupos + Playoff: divide los equipos en grupos y genera los partidos '
         'de la fase de grupos.')
sub_paso('Round Robin / Americano / Mexicano: genera todos los partidos entre '
         'todos los equipos.')

paso(4, 'Verificar el fixture generado',
     'El sistema redirige a la vista del fixture donde se puede ver el cuadro '
     'con todos los partidos. Si hay un número impar de equipos, algunos pueden '
     'recibir un "BYE" (pasan directamente a la siguiente ronda).')

atencion('El fixture solo puede generarse si la categoría tiene al menos el mínimo '
         'de equipos aprobados definido al crear la categoría. Si no se alcanza el mínimo, '
         'aparece el mensaje "Se necesitan al menos X equipos aprobados. Actualmente hay Y."')

separador()

h2('7.10 Cargar Resultados de Partidos')
body('Los resultados se cargan desde la vista de fixture de cada categoría.')

screenshot('Vista de fixture con partido seleccionado para cargar resultado')

paso(1, 'Navegar al fixture de la categoría',
     'Ir a Torneos → seleccionar el torneo → seleccionar la categoría → Fixture.')

paso(2, 'Seleccionar el partido',
     'En el cuadro de fixture, hacer clic sobre el partido cuyo resultado '
     'se desea cargar. Se abre un formulario o modal.')

paso(3, 'Ingresar el resultado por sets',
     'Para cada set completado, ingresar los games de cada equipo:')
campo('Games Equipo 1', 'Número de games del primer equipo en ese set.')
campo('Games Equipo 2', 'Número de games del segundo equipo en ese set.')
campo('Tiebreak Equipo 1', 'Puntos del tiebreak (solo si el set llegó al tiebreak).')
campo('Tiebreak Equipo 2', 'Puntos del tiebreak del segundo equipo.')
sub_paso('Ejemplo: para un partido que termina 6-4, 3-6, 7-5, ingresar tres sets '
         'con esos valores respectivamente.')

paso(4, 'Registrar walkover o abandono (si aplica)',
     'Si un equipo no se presenta o se retira durante el partido:')
sub_paso('Marcar la opción "Walkover" o "Abandono".')
sub_paso('Seleccionar el equipo que resultó ganador.')

paso(5, 'Confirmar el resultado',
     'Hacer clic en "Guardar resultado". El sistema determina automáticamente '
     'el ganador del partido y actualiza el cuadro.')
sub_paso('En eliminación directa: el ganador avanza automáticamente al siguiente partido.')
sub_paso('En grupos: la tabla de posiciones se actualiza inmediatamente.')

paso(6, 'Corregir un resultado ya cargado',
     'Si se cometió un error, hacer clic nuevamente sobre el partido en el fixture. '
     'El formulario muestra los valores actuales. Modificarlos y guardar. '
     'Cada modificación queda registrada en la auditoría del organizador.')

nota('Solo los usuarios con el permiso "Resultados" pueden cargar y modificar resultados.')

separador()

h2('7.11 Avanzar al Playoff (Formato Grupos + Playoff)')
body('Cuando se usa el formato Grupos + Playoff, al completar todos los partidos '
     'de la fase de grupos es necesario generar el cuadro de playoff manualmente.')

paso(1, 'Completar todos los partidos de grupos',
     'Cargar el resultado de cada partido de cada grupo. '
     'La tabla de posiciones se actualiza en tiempo real.')

paso(2, 'Verificar el aviso del sistema',
     'Cuando todos los partidos de grupos están completos, aparece una tarjeta verde '
     'con el mensaje "Fase de grupos completada — Todos los partidos de grupos terminaron".')

paso(3, 'Generar el playoff',
     'Hacer clic en el botón "→ Generar playoff" que aparece en esa tarjeta. '
     'El sistema clasifica automáticamente a los equipos según sus posiciones '
     'en cada grupo y genera el cuadro de eliminación del playoff.')

atencion('Si quedan partidos de grupos sin resultado, el botón "Generar playoff" '
         'no está disponible. El mensaje "Hay partidos de grupos sin completar" '
         'aparece si se intenta avanzar prematuramente.')

separador()

h2('7.12 Finalizar un Torneo')

paso(1, 'Verificar que todos los partidos estén completados',
     'Asegurarse de que el partido final de cada categoría tiene un resultado cargado.')

paso(2, 'Finalizar el torneo',
     'Desde la página de detalle del torneo, hacer clic en el botón "Finalizar torneo". '
     'El estado cambia a "COMPLETADO".')

paso(3, 'Actualización automática del ranking',
     'Si las categorías del torneo están vinculadas a una tabla de ranking, '
     'el sistema calcula automáticamente los puntos según las posiciones finales '
     'y actualiza los registros de ranking.')

paso(4, 'Revisar los reportes',
     'Con el torneo completado, la sección de reportes muestra el reporte de campeones '
     'con el ganador de cada categoría y los puntos otorgados al ranking.')

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  8. CALENDARIO
# ══════════════════════════════════════════════════════════════════════════════
h1('8. Panel del Organizador — Calendario')
body('El Calendario permite visualizar y gestionar todos los partidos programados '
     'de la organización en una vista mensual interactiva.')

h2('8.1 Navegar el Calendario   (/dashboard/calendario)')

screenshot('Vista mensual del calendario de partidos')

paso(1, 'Acceder al calendario',
     'Hacer clic en "Calendario" en el menú lateral izquierdo.')

paso(2, 'Entender el layout',
     'La pantalla está dividida en tres columnas:')
sub_paso('Columna izquierda — Calendario mini: muestra el mes actual con un mini-calendario. '
         'Los días con partidos tienen puntos de colores (verde: en curso, azul: programado, '
         'naranja: sin resultado, negro: jugado). Un botón "Hoy" vuelve al día actual.')
sub_paso('Columna central — Resumen del día seleccionado: muestra los contadores de partidos '
         'por estado y la lista de partidos del día con o sin horario asignado.')
sub_paso('Columna derecha / Tabla principal — Agenda del día: tabla completa con columnas '
         '"Hora | Cancha | Partido | Categoría | Estado".')

paso(3, 'Navegar entre meses',
     'Usar los botones "‹" y "›" a los costados del nombre del mes en el calendario mini '
     'para ir al mes anterior o siguiente.')

paso(4, 'Seleccionar un día',
     'Hacer clic en cualquier número del calendario mini. '
     'La columna central y la tabla principal se actualizan mostrando los partidos '
     'de ese día. El día seleccionado aparece resaltado en verde.')

paso(5, 'Filtrar por categoría',
     'Usar el filtro de categorías (chips o dropdown) en la parte superior '
     'para ver solo los partidos de una categoría específica.')

separador()

h2('8.2 Programar un Partido')
body('Los partidos sin fecha y hora asignadas aparecen en la sección '
     '"Sin horario" de la columna central.')

screenshot('Modal de programación de partido')

paso(1, 'Localizar los partidos sin programar',
     'En la columna central, buscar la sección "Sin horario" con su badge '
     'de cantidad en amarillo. Muestra los partidos pendientes de programar '
     'con el formato "Equipo1 vs Equipo2 · Torneo · Categoría".')

paso(2, 'Abrir el formulario de programación',
     'Hacer clic en un partido de la lista "Sin horario", o hacer clic en el botón '
     '"+ Agendar partido" del encabezado (que hace scroll a la sección). '
     'Se abre un modal o formulario con los campos de programación.')

paso(3, 'Completar los datos de programación', 'Ingresar la información requerida para agendar el partido:')
campo('Fecha', 'Selector de fecha para el partido.')
campo('Hora de inicio', 'Hora en formato HH:MM.')
campo('Hora de fin', 'Opcional. Hora estimada de finalización.')
campo('Sede', 'Seleccionar la sede donde se jugará el partido (de las sedes registradas).')
campo('Cancha', 'Seleccionar la cancha específica dentro de la sede elegida.')

paso(4, 'Guardar la programación',
     'Hacer clic en "Confirmar" o "Guardar". El partido pasa de la sección '
     '"Sin horario" a la agenda del día seleccionado. '
     'En la tabla principal aparece con su hora, cancha y estado.')

separador()

h2('8.3 Reprogramar un Partido')

paso(1, 'Seleccionar el día con el partido',
     'En el calendario mini, hacer clic en el día donde está agendado el partido.')

paso(2, 'Localizar el partido en la tabla',
     'En la tabla de agenda, encontrar el partido que se desea reprogramar.')

paso(3, 'Editar la programación',
     'Hacer clic en el partido para abrir el formulario de edición. '
     'Modificar la fecha, hora, sede o cancha según corresponda.')

paso(4, 'Guardar los cambios',
     'El sistema actualiza la programación y registra el cambio en el historial '
     'de reprogramaciones del partido.')

nota('El historial de reprogramaciones de un partido puede consultarse en la auditoría '
     'del organizador.')

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  9. JUGADORES
# ══════════════════════════════════════════════════════════════════════════════
h1('9. Panel del Organizador — Jugadores')
body('La sección de Jugadores permite buscar, ver y agregar jugadores al sistema.')

screenshot('Directorio de jugadores del organizador')

paso(1, 'Acceder a la sección',
     'Hacer clic en "Jugadores" en el menú lateral izquierdo.')

paso(2, 'Buscar un jugador',
     'Escribir el nombre, apellido, DNI o email en el campo de búsqueda. '
     'Los resultados se actualizan en tiempo real. '
     'Cada resultado muestra: avatar con iniciales, nombre completo (Apellido, Nombre), '
     'email, DNI y cantidad de equipos en los que participó.')

paso(3, 'Ver el detalle de un jugador',
     'Hacer clic en un jugador de la lista para ver su ficha completa: '
     'datos personales, historial de torneos e inscripciones.')

paso(4, 'Agregar un jugador manualmente',
     'Hacer clic en el botón "+ Agregar jugador". '
     'Se abre un formulario para ingresar los datos del jugador directamente: '
     'nombre, apellido, email, DNI, teléfono y fecha de nacimiento. '
     'Esto es útil para registrar jugadores que no tienen acceso al portal o '
     'para migrar datos de sistemas anteriores.')

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  10. SEDES Y CANCHAS
# ══════════════════════════════════════════════════════════════════════════════
h1('10. Panel del Organizador — Sedes y Canchas')
body('La sección de Sedes permite registrar los lugares donde se realizan los torneos '
     'y las canchas disponibles en cada uno.')

h2('10.1 Crear una Sede')

screenshot('Listado de sedes y formulario de nueva sede')

paso(1, 'Acceder a la sección de sedes',
     'Hacer clic en "Sedes" en el menú lateral, o acceder desde Configuración → "Sedes / Canchas".')

paso(2, 'Ver las sedes existentes',
     'La pantalla muestra una grilla de tarjetas, una por sede. '
     'Cada tarjeta indica: nombre, ciudad, dirección y cantidad de canchas.')

paso(3, 'Crear una nueva sede',
     'Hacer clic en el botón "+ Nueva sede". Completar el formulario:')
campo('Nombre', 'Nombre identificatorio de la sede (ej: "Club Pádel Norte"). Obligatorio.')
campo('Dirección', 'Dirección completa de la sede.')
campo('Ciudad', 'Ciudad donde se ubica.')
campo('URL de mapa', 'Enlace a Google Maps u otro servicio de mapas (opcional).')

paso(4, 'Guardar',
     'Hacer clic en "Guardar". La sede aparece en la grilla y ya está disponible '
     'para asignar al programar partidos.')

separador()

h2('10.2 Agregar Canchas a una Sede')

paso(1, 'Acceder al detalle de la sede',
     'Hacer clic en una tarjeta de sede en la grilla.')

paso(2, 'Ver las canchas existentes',
     'Se muestra una tabla con todas las canchas de esa sede: '
     'nombre, superficie, tipo (indoor/outdoor) y estado (activa/inactiva).')

paso(3, 'Agregar una nueva cancha',
     'Hacer clic en el botón "+ Agregar cancha". Completar el formulario:')
campo('Nombre', 'Nombre o número de la cancha (ej: "Cancha 1", "Central"). Obligatorio.')
campo('Superficie', 'Tipo de superficie: Cemento, Pasto sintético, Cristal, Césped u Otros.')
campo('Tipo', 'Indoor (cubierta) u Outdoor (descubierta).')

paso(4, 'Guardar',
     'Hacer clic en "Guardar". La cancha aparece en la tabla y está disponible '
     'para asignar al programar partidos.')

paso(5, 'Desactivar o reactivar una cancha',
     'Las canchas se pueden desactivar temporalmente (ej: por mantenimiento). '
     'Una cancha inactiva no aparece en el selector de canchas al programar partidos. '
     'Usar el toggle o el botón de acción en la fila de la cancha para cambiar su estado.')

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  11. RANKING
# ══════════════════════════════════════════════════════════════════════════════
h1('11. Panel del Organizador — Ranking')
body('La sección de Ranking permite llevar un registro histórico de puntos '
     'acumulados por jugadores y parejas a lo largo de torneos y temporadas.')

h2('11.1 Crear una Tabla de Ranking')

screenshot('Vista de ranking con tablas y reglas de puntuación')

paso(1, 'Acceder a la sección',
     'Hacer clic en "Ranking" en el menú lateral izquierdo.')

paso(2, 'Crear la tabla',
     'Hacer clic en el botón "+ Crear tabla". Completar el formulario:')
campo('Nombre', 'Nombre descriptivo (ej: "Ranking General 2026", "Ranking Mixtos"). Obligatorio.')
campo('Categoría', 'Categoría a la que aplica el ranking (opcional). '
      'Si se deja vacío, es un ranking general.')
campo('Temporada', 'Período al que corresponde (ej: "2026", "2026-Q1"). Obligatorio.')

paso(3, 'Confirmar',
     'Hacer clic en "Crear". La tabla aparece como una nueva pestaña en el selector '
     'de tablas de la parte superior de la pantalla.')

separador()

h2('11.2 Definir Reglas de Puntuación')
body('Cada tabla de ranking tiene sus propias reglas que determinan cuántos puntos '
     'recibe cada posición final. Si no se definen reglas, el sistema aplica valores '
     'predeterminados (1°: 100 pts, 2°: 60, 3°: 40, 4°: 20, participación: 10).')

paso(1, 'Seleccionar la tabla',
     'Hacer clic en la pestaña de la tabla de ranking que se desea configurar.')

paso(2, 'Ver las reglas existentes',
     'La sección "Reglas de puntuación" muestra cada regla con: '
     'posición, descripción y puntos en verde. '
     'Cada regla tiene botones "Editar" y "Eliminar".')

paso(3, 'Agregar una nueva regla',
     'En el formulario de la parte inferior de la sección, completar:')
campo('Posición', 'Selector con los valores: 1°, 2°, 3°, 4°, 5°, 6°, 7°, 8° y '
      '99 (participación). La posición 99 se reserva para otorgar puntos a todos '
      'los participantes que no queden en el podio.')
campo('Puntos', 'Cantidad de puntos a otorgar. Campo numérico obligatorio.')
campo('Descripción', 'Texto descriptivo opcional (ej: "Campeón", "Subcampeón").')
sub_paso('Hacer clic en "+ Agregar" para guardar la nueva regla.')

paso(4, 'Editar una regla existente',
     'Hacer clic en el botón "Editar" (ícono de lápiz) de la regla. '
     'Los campos se vuelven editables. Modificar los valores y hacer clic en "Guardar". '
     'Hacer clic en "Cancelar" para descartar los cambios.')

paso(5, 'Eliminar una regla',
     'Hacer clic en el botón "Eliminar" (ícono de papelera). '
     'El sistema elimina la regla inmediatamente.')

separador()

h2('11.3 Recalcular el Ranking')
body('Si se modifican reglas o resultados de partidos, el ranking puede recalcularse '
     'completamente desde cero.')

paso(1, 'Seleccionar la tabla',
     'Hacer clic en la pestaña de la tabla que se desea recalcular.')

paso(2, 'Iniciar el recálculo',
     'Hacer clic en el botón "Recalcular ranking". '
     'El sistema procesa todos los resultados históricos aplicando las reglas '
     'actuales y actualiza las posiciones y puntos de cada participante.')

paso(3, 'Verificar los resultados',
     'La tabla de posiciones se actualiza mostrando las nuevas posiciones, '
     'puntos acumulados y cantidad de torneos jugados por cada participante.')

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  12. REPORTES
# ══════════════════════════════════════════════════════════════════════════════
h1('12. Panel del Organizador — Reportes')
body('La sección de Reportes ofrece estadísticas consolidadas de la organización '
     'agrupadas en tres vistas: Inscriptos, Partidos y Campeones.')

screenshot('Pantalla de reportes con tabs y estadísticas')

paso(1, 'Acceder a los reportes',
     'Hacer clic en "Reportes" en el menú lateral izquierdo.')

paso(2, 'Ver el reporte de Inscriptos (pestaña "Inscriptos")',
     'Esta pestaña muestra:')
sub_paso('Cuatro tarjetas con totales: Total inscriptos, Aprobados, Pendientes y Rechazados.')
sub_paso('Tabla agrupada por torneo con columnas: Categoría, Total, Aprobados, '
         'Pendientes, Rechazados y Progreso (%).')
sub_paso('Fila de totales generales al final de la tabla.')
sub_paso('Botones "↓ PDF" y "↓ CSV" para descargar el reporte.')

paso(3, 'Ver el reporte de Partidos (pestaña "Partidos")',
     'Esta pestaña muestra:')
sub_paso('Cuatro tarjetas: Total partidos, Jugados, Pendientes y % Completado.')
sub_paso('Tabla con igual estructura que la de inscriptos pero con datos de partidos, '
         'agrupada por categoría y etapa.')

paso(4, 'Ver el reporte de Campeones (pestaña "Campeones")',
     'Esta pestaña muestra los resultados finales de todos los torneos completados:')
sub_paso('Si no hay torneos completados: ícono de trofeo con el mensaje '
         '"Sin torneos completados aún".')
sub_paso('Si hay torneos completados: tabla con columnas "Torneo | Categoría | '
         'Campeones | Formato | Fecha". Los campeones aparecen destacados con '
         'ícono de trofeo dorado.')

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  13. AUDITORÍA
# ══════════════════════════════════════════════════════════════════════════════
h1('13. Panel del Organizador — Auditoría')
body('La auditoría registra automáticamente todas las acciones realizadas en la organización. '
     'Permite rastrear quién hizo qué y cuándo.')

screenshot('Registro de auditoría del organizador')

paso(1, 'Acceder a la auditoría',
     'Hacer clic en "Auditoría" en el menú lateral izquierdo.')

paso(2, 'Explorar el registro de acciones',
     'Se muestra una lista de entradas ordenadas de la más reciente a la más antigua. '
     'Cada entrada muestra:')
sub_paso('Usuario que realizó la acción (nombre + email).')
sub_paso('Tipo de acción: Crear, Editar, Eliminar, Publicar, Aprobar, Rechazar, '
         'Cancelar, Cargar resultado, Modificar resultado.')
sub_paso('Entidad afectada: Torneo, Partido, Inscripción, Categoría, etc.')
sub_paso('Estado anterior y nuevo (datos completos del registro modificado).')
sub_paso('Fecha y hora exacta.')
sub_paso('Dirección IP del usuario al momento de la acción.')

paso(3, 'Filtrar el registro',
     'Aplicar los filtros disponibles:')
campo('Usuario', 'Ver solo las acciones de un usuario específico.')
campo('Tipo de acción', 'Filtrar por Crear, Editar, Eliminar, etc.')
campo('Entidad', 'Filtrar por tipo de entidad (Torneo, Partido, Inscripción, etc.).')
campo('Rango de fechas', 'Ver las acciones de un período específico.')

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  14. CONFIGURACIÓN
# ══════════════════════════════════════════════════════════════════════════════
h1('14. Configuración del Organizador')
body('La sección de Configuración está organizada en pestañas. Se accede desde '
     '"Configuración" en el menú lateral izquierdo.')

h2('14.1 Datos de la Organización (pestaña "Organización")')

screenshot('Configuración - Tab de datos de la organización')

paso(1, 'Acceder a la pestaña',
     'Hacer clic en "Configuración" en el menú lateral y luego en la pestaña "Organización".')

paso(2, 'Editar los datos',
     'Modificar los campos del formulario:')
campo('Nombre', 'Nombre público del organizador.')
campo('Slug', 'Identificador único para la URL pública (ej: "club-padel-norte"). '
      'Solo letras minúsculas, números y guiones.')
campo('Descripción', 'Texto descriptivo de la organización.')
campo('URL del logo', 'Enlace a la imagen del logo.')
campo('Email de contacto', 'Email público de la organización.')
campo('Teléfono', 'Teléfono de contacto.')
campo('Sitio web', 'URL del sitio web de la organización.')

paso(3, 'Guardar los cambios',
     'Hacer clic en "Guardar". El sistema confirma los cambios con un mensaje de éxito.')

separador()

h2('14.2 Parámetros Generales (pestaña "Parámetros")')
body('Define los valores que se usan por defecto al crear nuevos torneos. '
     'Pueden sobreescribirse individualmente en cada torneo o categoría.')

paso(1, 'Acceder a la pestaña "Parámetros"', 'Hacer clic en "Configuración" en el menú lateral y luego seleccionar la pestaña "Parámetros".')

paso(2, 'Configurar los valores default', 'Editar los siguientes campos y guardar los cambios al finalizar:')
campo('Sets por partido', 'Número de sets (por defecto: 3).')
campo('Games por set', 'Número de games por set (por defecto: 6).')
campo('Máximo de equipos por categoría', 'Cupo máximo default (por defecto: 16).')
campo('¿Permitir inscripción pública?', 'Casilla que habilita o deshabilita la inscripción '
      'desde el portal público.')
campo('¿Requerir pago?', 'Si está habilitado, las inscripciones solo se aprueban '
      'al confirmar el pago.')
campo('Zona horaria', 'Zona horaria de la organización para mostrar horarios correctamente.')

paso(3, 'Guardar los cambios',
     'Hacer clic en "Guardar". Los nuevos valores se aplican en los próximos torneos creados.')

separador()

h2('14.3 Gestión de Categorías (pestaña "Categorías")')
body('Define el catálogo de categorías disponibles para crear en los torneos.')

screenshot('Configuración - Tab de categorías')

paso(1, 'Ver las categorías existentes',
     'La tabla muestra todas las categorías con: nombre, género, nivel y estado '
     '(activa/inactiva).')

paso(2, 'Crear una nueva categoría',
     'Hacer clic en "+ Nueva categoría". Completar el formulario:')
campo('Nombre', 'Nombre de la categoría (ej: "1ra Masculina"). Debe ser único en la organización.')
campo('Género', 'Masculino, Femenino, Mixto u Abierto.')
campo('Nivel', 'Texto libre para indicar el nivel (ej: "1ra", "2da", "Principiantes").')

paso(3, 'Guardar',
     'Hacer clic en "Guardar". La categoría queda disponible para usarse en nuevos torneos.')

paso(4, 'Desactivar una categoría',
     'Usar el toggle o el botón de acción en la fila de la categoría para desactivarla. '
     'Las categorías inactivas no aparecen al crear nuevos torneos, '
     'pero las existentes en torneos ya creados no se ven afectadas.')

atencion('No es posible eliminar una categoría que ya esté asociada a uno o más torneos. '
         'Desactivarla es la opción correcta en ese caso.')

separador()

h2('14.4 Gestión de Colaboradores (pestaña "Colaboradores")')
body('Permite invitar y administrar a los miembros del equipo organizador.')

screenshot('Configuración - Tab de colaboradores')

paso(1, 'Ver los miembros actuales',
     'La lista muestra todos los miembros con: nombre, badge de rol '
     '(Propietario en amarillo, Organizador en azul, Colaborador en gris), '
     'email y permisos asignados. El miembro actual aparece con el badge "Vos" en verde.')

paso(2, 'Invitar a un nuevo colaborador',
     'En el campo de invitación en la parte superior:')
sub_paso('Escribir el email del usuario que se desea invitar.')
sub_paso('Hacer clic en "Invitar". El sistema valida que el email corresponda '
         'a una cuenta registrada.')
sub_paso('Si el email no existe en el sistema, aparece un mensaje de error en rojo. '
         'El usuario debe crear su cuenta primero.')
sub_paso('Si el email es válido, el usuario se agrega como Colaborador con permisos básicos.')

paso(3, 'Asignar permisos a un Colaborador',
     'En la tarjeta del colaborador, a la derecha, aparecen los checkboxes de permisos:')
sub_paso('Torneos, Inscripciones, Resultados, Calendario, Reportes.')
sub_paso('Marcar o desmarcar los permisos deseados.')
sub_paso('También se puede configurar el acceso a torneos específicos: '
         'activar "Solo torneos seleccionados" y marcar los torneos a los que '
         'el colaborador puede acceder.')
sub_paso('Hacer clic en "Guardar acceso" para confirmar los cambios.')

paso(4, 'Cambiar el rol de un miembro',
     'Hacer clic en el badge de rol del miembro para ver las opciones disponibles '
     'y seleccionar el nuevo rol.')

paso(5, 'Eliminar a un colaborador',
     'Hacer clic en el botón "Eliminar" (rojo) en la tarjeta del miembro. '
     'El sistema solicita confirmación antes de eliminar.')

nota('Solo los usuarios con rol Owner pueden gestionar colaboradores. '
     'Un Owner no puede eliminarse a sí mismo ni cambiar su propio rol.')

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  15. ADMINISTRACIÓN GLOBAL
# ══════════════════════════════════════════════════════════════════════════════
h1('15. Panel de Administración Global')
body('El Panel de Administración Global está disponible exclusivamente para usuarios '
     'con el rol Super Admin. Se accede desde el menú de usuario del encabezado → '
     '"Panel admin →".')

h2('15.1 Dashboard de Administración   (/admin)')

screenshot('Dashboard de administración global')

paso(1, 'Ver las estadísticas globales',
     'La pantalla muestra cuatro tarjetas con:')
sub_paso('Organizadores activos / Total en el sistema.')
sub_paso('Usuarios registrados en total.')
sub_paso('Torneos totales en todos los organizadores.')
sub_paso('Eventos registrados en los últimos 7 días.')

paso(2, 'Revisar la actividad reciente',
     'La sección inferior lista las últimas 20 acciones realizadas en todo el sistema: '
     'quién hizo qué, en qué organizador y cuándo. '
     'El enlace "Ver todo auditoría" lleva al registro completo.')

separador()

h2('15.2 Gestión de Usuarios   (/admin/usuarios)')

screenshot('Gestión global de usuarios')

paso(1, 'Ver la tabla de usuarios',
     'La tabla lista todos los usuarios registrados en el sistema con: '
     'nombre, email, rol de sistema (Jugador o Super Admin), '
     'organizadores a los que pertenece, estado y fecha de creación.')

paso(2, 'Activar o desactivar una cuenta',
     'Hacer clic en el toggle de estado en la fila del usuario. '
     'Un usuario desactivado no puede iniciar sesión.')

paso(3, 'Promover a Super Admin',
     'Hacer clic en el botón correspondiente en la fila del usuario. '
     'El sistema solicita confirmación. Al promover, el usuario puede acceder '
     'al panel de administración global.')

paso(4, 'Revocar el rol de Super Admin',
     'Hacer clic en el botón "Revocar Admin" en la fila del usuario Super Admin. '
     'El usuario vuelve al rol de Jugador.')

atencion('Un Super Admin no puede modificar su propio rol ni desactivarse a sí mismo '
         'desde esta pantalla.')

separador()

h2('15.3 Gestión de Organizadores   (/admin/organizadores)')

paso(1, 'Ver la tabla de organizadores',
     'Se listan todos los organizadores con: nombre, slug, cantidad de miembros, '
     'cantidad de torneos y estado (activo/inactivo).')

paso(2, 'Editar un organizador',
     'Hacer clic en el botón de edición de la fila. '
     'Permite modificar el nombre, slug y otros datos del organizador.')

paso(3, 'Ver los miembros de un organizador',
     'Hacer clic en el botón "Ver miembros" para ver el listado de usuarios '
     'que pertenecen a ese organizador.')

paso(4, 'Activar o desactivar un organizador',
     'Hacer clic en el toggle de estado. Un organizador desactivado no puede '
     'ser accedido por sus miembros.')

separador()

h2('15.4 Auditoría Global   (/admin/auditoria)')
body('Similar a la auditoría del organizador, pero con alcance a todo el sistema.')

paso(1, 'Ver el registro global',
     'Se muestran todas las acciones de todos los organizadores, usuarios y torneos '
     'del sistema, con: usuario, organizador, tipo de acción, entidad afectada, '
     'datos antes/después, fecha y hora.')

paso(2, 'Aplicar filtros',
     'Los filtros disponibles incluyen: organizador, usuario, tipo de acción, '
     'entidad y rango de fechas. Usar los filtros combinados para encontrar '
     'eventos específicos en el historial global.')

screenshot('Auditoría global del sistema')

# ── Pie de documento ──────────────────────────────────────────────────────────
separador()
pie = doc.add_paragraph()
pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
pie.paragraph_format.space_before = Pt(16)
rp = pie.add_run(
    'Fin del Manual de Usuario  ·  Sistema de Gestión de Torneos de Pádel  ·  v2.0'
)
rp.font.size      = Pt(10)
rp.font.color.rgb = GRIS_CLARO
rp.italic         = True

# ── Guardar ───────────────────────────────────────────────────────────────────
out = r'C:\proyectos\torneos-padel\Manual_de_Usuario_v2_PasoAPaso.docx'
doc.save(out)
print(f'Manual generado en: {out}')

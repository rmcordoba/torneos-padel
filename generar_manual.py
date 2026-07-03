"""
Genera el Manual de Usuario del Sistema de Gestión de Torneos de Padel en formato .docx
Ejecutar: python generar_manual.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ─── Pagina ───────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width   = Cm(21)
section.page_height  = Cm(29.7)
section.left_margin  = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin   = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ─── Colores ──────────────────────────────────────────────────────────
C_DARK   = RGBColor(0x1E, 0x3A, 0x5F)
C_BLUE   = RGBColor(0x1A, 0x56, 0xDB)
C_GREEN  = RGBColor(0x22, 0xC5, 0x5E)
C_GRAY   = RGBColor(0x6B, 0x72, 0x80)
C_ORANGE = RGBColor(0xB4, 0x5A, 0x09)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

# ─── Estilos de heading ───────────────────────────────────────────────
styles = doc.styles

def _style(name, size, bold, color, before=12, after=4):
    s = styles[name]
    s.font.name  = 'Calibri'
    s.font.size  = Pt(size)
    s.font.bold  = bold
    s.font.color.rgb = color
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after  = Pt(after)

_style('Heading 1', 16, True,  C_DARK,  18, 6)
_style('Heading 2', 13, True,  C_BLUE,  12, 4)
_style('Heading 3', 11.5, True, C_BLUE, 10, 3)
_style('Heading 4', 11, True,  C_DARK,   8, 2)

n = styles['Normal']
n.font.name = 'Calibri'
n.font.size = Pt(11)
n.paragraph_format.space_after = Pt(6)

# ─── Helpers ──────────────────────────────────────────────────────────

def pb():
    doc.add_page_break()

def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def h3(t): doc.add_heading(t, level=3)
def h4(t): doc.add_heading(t, level=4)
def body(t): doc.add_paragraph(t)

def bullet(text, lvl=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(lvl * 0.5 + 0.5)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(11)

def step(text):
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(11)

def note(text):
    p = doc.add_paragraph()
    r = p.add_run(f"NOTA: {text}")
    r.font.color.rgb = C_DARK
    r.font.size = Pt(10)
    r.font.italic = True
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.right_indent = Cm(1)

def warn(text):
    p = doc.add_paragraph()
    r = p.add_run(f"IMPORTANTE: {text}")
    r.font.color.rgb = C_ORANGE
    r.font.size = Pt(10)
    r.font.bold  = True
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.right_indent = Cm(1)

def shot(caption="[CAPTURA DE PANTALLA]"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"\n{'─'*55}\n{caption}\n{'─'*55}\n")
    r.font.color.rgb = C_GRAY
    r.font.size = Pt(9)
    r.font.italic = True

def _shd(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    tcPr.append(shd)

def table_hdr(table, cols, fill='1A56DB'):
    row = table.rows[0]
    for i, col in enumerate(cols):
        cell = row.cells[i]
        cell.text = col
        r = cell.paragraphs[0].runs[0]
        r.bold = True
        r.font.color.rgb = C_WHITE
        _shd(cell, fill)

def toc():
    p = doc.add_paragraph()
    r = p.add_run()
    def fld(tag): e = OxmlElement(tag); return e

    fc_begin = OxmlElement('w:fldChar')
    fc_begin.set(qn('w:fldCharType'), 'begin')
    r._r.append(fc_begin)

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-4" \\h \\z \\u '
    r._r.append(instr)

    fc_sep = OxmlElement('w:fldChar')
    fc_sep.set(qn('w:fldCharType'), 'separate')
    r._r.append(fc_sep)

    t = OxmlElement('w:t')
    t.text = 'Haga clic aqui y luego: clic derecho -> Actualizar campo -> Actualizar toda la tabla'
    r._r.append(t)

    fc_end = OxmlElement('w:fldChar')
    fc_end.set(qn('w:fldCharType'), 'end')
    r._r.append(fc_end)


# ══════════════════════════════════════════════════════════════
#  PORTADA
# ══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('\n\n\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('SISTEMA DE GESTION DE TORNEOS DE PADEL')
r.font.name='Calibri'; r.font.size=Pt(26); r.font.bold=True; r.font.color.rgb=C_DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PadelPro')
r.font.name='Calibri'; r.font.size=Pt(20); r.font.color.rgb=C_BLUE

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('MANUAL DE USUARIO')
r.font.name='Calibri'; r.font.size=Pt(22); r.font.bold=True; r.font.color.rgb=C_GREEN

doc.add_paragraph()
shot('[ LOGO / IMAGEN INSTITUCIONAL ]')
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Version 1.0  -  2026')
r.font.name='Calibri'; r.font.size=Pt(12); r.font.color.rgb=C_GRAY

pb()

# ══════════════════════════════════════════════════════════════
#  INDICE
# ══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('INDICE DE CONTENIDOS')
r.font.name='Calibri'; r.font.size=Pt(16); r.font.bold=True; r.font.color.rgb=C_DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    'Para actualizar la numeracion de paginas: clic derecho sobre el indice '
    '-> "Actualizar campo" -> "Actualizar toda la tabla".'
)
r.font.size=Pt(9); r.font.italic=True; r.font.color.rgb=C_GRAY

doc.add_paragraph()
toc()
pb()


# ══════════════════════════════════════════════════════════════
#  SECCION INICIAL
# ══════════════════════════════════════════════════════════════
h1('MANUAL DE USUARIO - SECCION INICIAL')

# ── 1. INTRODUCCION ───────────────────────────────────────────
h1('1. INTRODUCCION')
body(
    'El presente manual de usuario tiene como objetivo brindar una guia clara, ordenada y accesible '
    'para el uso del Sistema de Gestion de Torneos de Padel - PadelPro. '
    'Este sistema fue desarrollado para digitalizar y optimizar la organizacion de torneos de padel, '
    'permitiendo gestionar inscripciones, fixtures, resultados, calendarios, rankings y reportes '
    'desde una plataforma web multi-tenant, accesible desde cualquier dispositivo con conexion a Internet.'
)
body(
    'A traves de este manual, el usuario podra comprender el funcionamiento de cada modulo, '
    'identificar las acciones disponibles segun su rol y realizar todas las operaciones necesarias '
    'de manera segura y eficiente, sin necesidad de conocimientos tecnicos avanzados.'
)

h2('1.1 Proposito del Manual')
body('El proposito de este manual es:')
bullet('Explicar detalladamente como navegar y utilizar cada seccion del sistema.')
bullet('Ofrecer instrucciones paso a paso para la correcta realizacion de cada operacion.')
bullet('Facilitar el aprendizaje de nuevos usuarios y servir como material de consulta ante dudas operativas.')
bullet('Garantizar que el uso del sistema se realice siguiendo las buenas practicas establecidas.')
body('Este documento actua como guia oficial de uso del sistema para todo el personal autorizado.')

h2('1.2 Audiencia')
body('Este manual esta dirigido a:')
bullet('Super Administradores del Sistema: responsables de la configuracion global, gestion de organizadores y usuarios, y auditoria del sistema.')
bullet('Organizadores: encargados de crear y gestionar torneos, sedes, canchas, categorias, inscripciones y resultados.')
bullet('Colaboradores: usuarios con permisos limitados dentro de una organizacion (inscripciones, resultados, agenda).')
bullet('Jugadores: participantes que gestionan su perfil e inscripciones a torneos.')
bullet('Espectadores: visitantes sin cuenta que acceden al portal publico para consultar torneos, cuadros, agenda y rankings.')
body('No se requieren conocimientos tecnicos previos. Las instrucciones estan disenadas para usuarios con distintos niveles de experiencia.')

pb()

# ── 2. REQUISITOS ─────────────────────────────────────────────
h1('2. REQUISITOS PARA ACCEDER AL SISTEMA')
body('Antes de ingresar al sistema, el usuario debe cumplir con los siguientes requisitos:')
bullet('Poseer un usuario y contrasena asignados, o haberse registrado a traves del formulario de alta.')
bullet('Contar con conexion a Internet estable.')
bullet('Disponer de un dispositivo compatible (PC, notebook o tablet) con un navegador actualizado.')
bullet('Ingresar desde un entorno seguro para evitar el acceso de terceros.')

h2('2.1 Dispositivos y Navegadores Compatibles')
body('El sistema es compatible con los siguientes entornos:')
bullet('Google Chrome 110 o superior (recomendado).')
bullet('Mozilla Firefox 110 o superior.')
bullet('Microsoft Edge 110 o superior.')
bullet('Safari 16 o superior (macOS / iOS).')
bullet('Resolucion minima recomendada: 1280 x 720 pixeles.')
note('El sistema es completamente responsivo y puede utilizarse desde dispositivos moviles, aunque se recomienda el uso en pantallas de mayor tamano para una mejor experiencia de usuario.')

h2('2.2 Acceso al Sistema')
body(
    'Para acceder al sistema, el usuario debe ingresar la direccion web (URL) correspondiente '
    'en el navegador. La pantalla inicial mostrara directamente la pagina de inicio de sesion.'
)

pb()

# ── 3. LOGIN ──────────────────────────────────────────────────
h1('3. PANTALLA DE INICIO DE SESION / LOGIN')
shot('[ CAPTURA DE PANTALLA - Pantalla de inicio de sesion ]')

h2('3.1 Descripcion General')
body(
    'La pantalla de Inicio de Sesion permite al usuario autenticarse en el sistema mediante el '
    'ingreso de sus credenciales personales. Esta funcionalidad garantiza que unicamente usuarios '
    'registrados y autorizados puedan acceder a las distintas secciones, asegurando la '
    'confidencialidad e integridad de la informacion. '
    'El acceso esta protegido mediante autenticacion basada en correo electronico y contrasena.'
)

h2('3.2 Elementos de la Interfaz')
body('La pantalla esta compuesta por los siguientes componentes:')
bullet('Campo "Correo electronico": permite ingresar el email registrado. Obligatorio.')
bullet('Campo "Contrasena": permite ingresar la contrasena asociada al usuario. Incluye icono para mostrar/ocultar el texto.')
bullet('Boton "Ingresar": ejecuta la validacion de credenciales. Si son correctas, redirige al panel segun el rol.')
bullet('Enlace "Olvidaste tu contrasena?": inicia el proceso de recuperacion de acceso.')
bullet('Enlace "Crear cuenta": redirige al formulario de registro.')

h2('3.3 Procedimiento de Inicio de Sesion')
body('Para acceder al sistema:')
step('Ingresar el correo electronico en el campo correspondiente.')
step('Ingresar la contrasena.')
step('Presionar el boton "Ingresar".')
step('Aguardar la validacion automatica del sistema.')
body('Si las credenciales son validas, el sistema habilitara el acceso al panel principal segun el rol asignado.')

h3('3.3.1 Validaciones y Mensajes')
body('El sistema realiza las siguientes validaciones:')
bullet('Verificacion de campos obligatorios.')
bullet('Validacion del formato del correo electronico.')
bullet('Confirmacion de existencia del usuario.')
bullet('Verificacion de coincidencia de contrasena.')
body('En caso de error se mostrara un mensaje informativo indicando la causa.')
shot('[ CAPTURA DE PANTALLA - Mensaje de error en login ]')

h2('3.4 Registro de Usuario')
h3('3.4.1 Descripcion General')
body(
    'La pantalla de Registro permite crear una nueva cuenta en el sistema. '
    'Posibilita la incorporacion de nuevos usuarios que deseen participar como jugadores '
    'o gestionar torneos como organizadores.'
)
shot('[ CAPTURA DE PANTALLA - Formulario de registro ]')

h3('3.4.2 Elementos de la Interfaz')
body('El formulario contiene los siguientes campos:')
bullet('Campo "Nombre": nombre del usuario. Obligatorio.')
bullet('Campo "Apellido": apellido del usuario. Obligatorio.')
bullet('Campo "Correo electronico": direccion valida y no registrada previamente. Obligatorio.')
bullet('Campo "Contrasena": contrasena de acceso. Obligatorio.')
bullet('Campo "Repetir contrasena": confirmacion de la contrasena. Obligatorio.')
bullet('Boton "Registrarme": valida y guarda los datos.')

h3('3.4.3 Procedimiento de Registro')
body('Para crear una nueva cuenta:')
step('Ingresar nombre y apellido.')
step('Ingresar un correo electronico valido.')
step('Definir una contrasena segura.')
step('Repetir la contrasena para confirmacion.')
step('Presionar el boton "Registrarme".')
body('Si los datos son correctos, el sistema registrara al usuario y mostrara una notificacion de exito. Luego podra iniciar sesion con sus credenciales.')

h2('3.5 Recuperar Contrasena')
body('En caso de no recordar la contrasena, hacer clic en el enlace "Olvidaste tu contrasena?" en la pantalla de inicio de sesion.')
shot('[ CAPTURA DE PANTALLA - Pantalla de recuperacion de contrasena ]')
body('El procedimiento es el siguiente:')
step('Hacer clic en "Olvidaste tu contrasena?".')
step('Ingresar el correo electronico registrado.')
step('Presionar "Enviar enlace de recuperacion".')
step('Revisar el correo y hacer clic en el enlace recibido.')
step('Ingresar la nueva contrasena y confirmarla.')
step('Presionar "Confirmar" para guardar los cambios.')
note('El enlace de recuperacion tiene una validez de 1 hora. Vencido ese plazo debera solicitarse uno nuevo.')

pb()

# ══════════════════════════════════════════════════════════════
#  4. PORTAL PUBLICO
# ══════════════════════════════════════════════════════════════
h1('4. PORTAL PUBLICO (Acceso sin Cuenta)')

h2('4.1 Descripcion General')
body(
    'El Portal Publico es el area del sistema accesible para cualquier persona sin necesidad de '
    'registrarse ni iniciar sesion. Permite consultar informacion actualizada sobre torneos, '
    'cuadros de competencia, agenda de partidos, rankings y directorio de jugadores. '
    'La navegacion se realiza a traves de un menu superior con pestanas.'
)
shot('[ CAPTURA DE PANTALLA - Portal publico, pantalla principal ]')

h2('4.2 Seccion Torneos')
body(
    'Muestra todos los torneos publicados. El torneo destacado (activo o mas proximo) '
    'se presenta en un banner principal con informacion completa.'
)
body('Informacion disponible por torneo:')
bullet('Nombre y descripcion.')
bullet('Fechas de inicio y fin.')
bullet('Estado (Inscripciones abiertas, En curso, Finalizado, etc.).')
bullet('Categorias disponibles con cupos, precio y estado de inscripcion.')
bullet('Sede y ubicacion.')
shot('[ CAPTURA DE PANTALLA - Listado de torneos publicos ]')

h2('4.3 Seccion Cuadros (Fixture)')
body(
    'Permite visualizar el cuadro de competencia del torneo activo o seleccionado, '
    'con selector de categoria. Muestra el estado de cada partido y los resultados registrados.'
)
body('Formatos disponibles para visualizar:')
bullet('Eliminacion simple (bracket).')
bullet('Grupos + Playoff.')
bullet('Doble eliminacion (cuadro A y cuadro B).')
bullet('Round Robin / Liga.')
bullet('Americano.')
bullet('Mexicano.')
shot('[ CAPTURA DE PANTALLA - Cuadro de competencia publico ]')

h2('4.4 Seccion Agenda')
body(
    'Muestra el calendario de partidos programados del torneo activo, organizado por dia. '
    'El usuario puede navegar entre fechas y consultar horarios, canchas y equipos participantes.'
)
shot('[ CAPTURA DE PANTALLA - Agenda publica ]')

h2('4.5 Seccion Ranking')
body(
    'Presenta las tablas de ranking con podio destacado para los tres primeros lugares '
    'y listado completo de posiciones con puntos acumulados y torneos jugados.'
)
shot('[ CAPTURA DE PANTALLA - Tabla de ranking publica ]')

h2('4.6 Seccion Jugadores')
body(
    'Directorio publico de jugadores registrados. '
    'Permite buscar por nombre y acceder a un perfil resumido con historial de torneos y estadisticas.'
)
shot('[ CAPTURA DE PANTALLA - Directorio de jugadores ]')

pb()

# ══════════════════════════════════════════════════════════════
#  5. PANEL DE JUGADOR
# ══════════════════════════════════════════════════════════════
h1('5. PANEL DE JUGADOR')

h2('5.1 Descripcion General')
body(
    'El Panel de Jugador es el area privada a la que acceden los usuarios con rol Jugador. '
    'Desde aqui pueden consultar torneos disponibles, gestionar inscripciones y actualizar su perfil.'
)
body('Menu lateral disponible:')
bullet('Mi Panel.')
bullet('Mi Perfil.')
bullet('Portal Publico.')
shot('[ CAPTURA DE PANTALLA - Dashboard del jugador ]')

h2('5.2 Mi Panel (Dashboard del Jugador)')
body('Pantalla principal del jugador. Muestra:')
bullet('Torneos abiertos a inscripcion con categorias disponibles, cupos y precios.')
bullet('Mis inscripciones activas (pareja, torneo, categoria, estado de aprobacion).')
bullet('Acceso rapido al portal publico.')

h2('5.3 Inscripcion a un Torneo')
body('Para inscribirse en un torneo:')
step('Identificar el torneo y categoria deseada en la seccion de torneos disponibles.')
step('Hacer clic en "Inscribirme".')
step('Seleccionar o ingresar el nombre del companero/a de pareja.')
step('Confirmar la inscripcion.')
body(
    'La inscripcion quedara en estado Pendiente hasta que el organizador la apruebe o rechace. '
    'El jugador recibira notificacion por correo electronico con la resolucion.'
)
shot('[ CAPTURA DE PANTALLA - Formulario de inscripcion a torneo ]')
warn(
    'Las inscripciones son siempre por pareja (dupla fija). Si el cupo esta completo, '
    'la inscripcion ingresa automaticamente a la lista de espera.'
)

h2('5.4 Mis Inscripciones')
body('Historial completo de inscripciones:')
bullet('Torneo y categoria.')
bullet('Companero/a de pareja.')
bullet('Fecha de inscripcion.')
bullet('Estado: Pendiente / Aprobada / Rechazada / En lista de espera.')
shot('[ CAPTURA DE PANTALLA - Listado de inscripciones del jugador ]')

h2('5.5 Mi Perfil')
body('Permite consultar y actualizar la informacion personal:')
bullet('Nombre y apellido.')
bullet('Correo electronico (no modificable).')
bullet('Telefono de contacto.')
bullet('Fecha de nacimiento.')
bullet('Numero de documento (DNI).')
bullet('Opcion para cambiar contrasena.')
shot('[ CAPTURA DE PANTALLA - Perfil del jugador ]')

h3('5.5.1 Cambio de Contrasena')
step('Acceder a "Mi Perfil" desde el menu lateral.')
step('Hacer clic en "Cambiar contrasena".')
step('Ingresar la contrasena actual.')
step('Ingresar la nueva contrasena y confirmarla.')
step('Guardar los cambios.')

pb()

# ══════════════════════════════════════════════════════════════
#  6. MODULO TORNEOS
# ══════════════════════════════════════════════════════════════
h1('6. MODULO DE TORNEOS (Organizador)')

h2('6.1 Descripcion General')
body(
    'El Modulo de Torneos es el nucleo operativo para los organizadores. '
    'Permite crear y gestionar torneos completos: desde la configuracion inicial hasta '
    'el cierre y generacion de resultados. Cada torneo puede contener multiples categorias '
    'con formatos de competencia independientes.'
)
body('Acceso: menu lateral -> Torneos.')

h2('6.2 Listado de Torneos')
body('Muestra todos los torneos del organizador organizados por estado del ciclo de vida:')
bullet('Activos / En curso: torneos publicados o en competencia.')
bullet('Proximos: inscripciones abiertas o cerradas, aun no iniciados.')
bullet('Finalizados: torneos completados o cancelados.')
body('Filtro disponible por estado: Borrador, Publicado, Inscripciones abiertas, En curso, Finalizado, Cancelado.')
body('Cada tarjeta muestra: nombre, estado, fechas, cantidad de categorias y equipos.')
shot('[ CAPTURA DE PANTALLA - Listado de torneos ]')

h2('6.3 Crear Nuevo Torneo')
body('Hacer clic en "Nuevo torneo" en la esquina superior derecha.')
shot('[ CAPTURA DE PANTALLA - Formulario de nuevo torneo ]')

h3('6.3.1 Datos Generales del Torneo')
bullet('Campo "Nombre": nombre del torneo. Obligatorio.')
bullet('Campo "Descripcion": descripcion del evento.')
bullet('Campo "Fecha de inicio" y "Fecha de fin". Obligatorios.')
bullet('Campo "Fecha limite de inscripcion": hasta cuando pueden inscribirse las parejas.')
bullet('Sede: seleccion de la sede donde se realizara el torneo.')
bullet('Torneo publico: si esta activado, aparece en el Portal Publico.')

h3('6.3.2 Categorias del Torneo')
body('Cada torneo puede tener multiples categorias. Campos por categoria:')
bullet('Categoria: seleccion del catalogo del organizador (ej. 4ta, 5ta, Mixto).')
bullet('Formato de competencia: define el tipo de cuadro.')
bullet('Maximo y minimo de parejas.')
bullet('Sets por partido y games por set.')
bullet('Precio por pareja.')
bullet('Tamano de grupo y equipos que avanzan por grupo (si aplica).')

h3('6.3.3 Formatos de Competencia Disponibles')
bullet('Eliminacion Simple: cuadro knockout. El perdedor queda eliminado.')
bullet('Grupos + Playoff: fase de grupos seguida de cuadro eliminatorio con los clasificados.')
bullet('Doble Eliminacion: dos cuadros paralelos. Gran Final entre ganadores de A y B.')
bullet('Round Robin (Liga): todos contra todos. Clasificacion por puntos.')
bullet('Americano: todos contra todos. Clasificacion por games ganados.')
bullet('Mexicano: rondas dinamicas con emparejamiento por ranking acumulado. Rondas configurables.')
note('Una vez generado el fixture, el formato y algunos parametros no pueden modificarse.')

body('Procedimiento de creacion:')
step('Completar los datos generales del torneo.')
step('Hacer clic en "Agregar categoria" y configurar cada una.')
step('Presionar "Guardar torneo".')
body('El torneo se creara en estado Borrador.')

h2('6.4 Detalle y Gestion del Torneo')
body('Al hacer clic en un torneo se accede a la pantalla de detalle con todas las acciones de gestion.')
shot('[ CAPTURA DE PANTALLA - Detalle del torneo ]')

h3('6.4.1 Transiciones de Estado del Torneo')
body('Los torneos siguen un ciclo de vida definido. Los botones de accion cambian segun el estado actual:')

t = doc.add_table(rows=8, cols=3)
t.style = 'Table Grid'
table_hdr(t, ['Estado Actual', 'Accion disponible', 'Nuevo Estado'])
data = [
    ('Borrador',                  'Publicar',               'Publicado'),
    ('Publicado',                 'Abrir inscripciones',     'Inscripciones abiertas'),
    ('Inscripciones abiertas',    'Cerrar inscripciones',    'Inscripciones cerradas'),
    ('Inscripciones cerradas',    'Reabrir inscripciones',   'Inscripciones abiertas'),
    ('Inscripciones cerradas',    'Iniciar torneo',          'En curso'),
    ('En curso',                  'Completar torneo',        'Finalizado'),
    ('Cualquier estado activo',   'Cancelar torneo',         'Cancelado'),
]
for i,(a,b,c) in enumerate(data):
    t.rows[i+1].cells[0].text = a
    t.rows[i+1].cells[1].text = b
    t.rows[i+1].cells[2].text = c

doc.add_paragraph()
warn('La accion de cancelar un torneo es irreversible. El sistema solicitara confirmacion antes de ejecutarla.')

h3('6.4.2 Editar Torneo')
body(
    'Permite modificar los datos generales (nombre, descripcion, fechas, sede, visibilidad). '
    'Disponible mientras el torneo no haya finalizado.'
)

h3('6.4.3 Gestion de Categorias en el Torneo')
body('La pantalla de detalle muestra una grilla con todas las categorias. Cada tarjeta permite:')
bullet('Ver el estado de la categoria (Borrador, Abierta, Cerrada).')
bullet('Acceder a las inscripciones de esa categoria.')
bullet('Acceder al fixture.')
bullet('Acceder al calendario de partidos.')
bullet('Editar los parametros de la categoria.')
bullet('Eliminar la categoria (solo si no tiene fixture generado).')
bullet('Cambiar el estado: Borrador -> Abierta -> Cerrada.')
shot('[ CAPTURA DE PANTALLA - Grilla de categorias del torneo ]')

h2('6.5 Inscripciones por Categoria')
body('Al acceder a las inscripciones de una categoria, se visualizan las parejas agrupadas por estado.')
shot('[ CAPTURA DE PANTALLA - Gestion de inscripciones por categoria ]')
body('Estados de inscripcion:')
bullet('Pendiente: esperando aprobacion del organizador.')
bullet('Aprobada: pareja confirmada en el torneo.')
bullet('Rechazada: pareja no aceptada.')
bullet('En lista de espera: cupo lleno, pareja en cola automatica.')
body('Acciones disponibles:')
bullet('Aprobar inscripcion individual: confirma la pareja. El jugador recibe notificacion por email.')
bullet('Rechazar inscripcion individual: declina la solicitud. El jugador recibe notificacion por email.')
bullet('"Aprobar todas las pendientes": aparece cuando hay mas de una pendiente. Aprueba todas de una vez.')
bullet('Registrar pareja manualmente: el organizador inscribe una pareja directamente.')
bullet('Lista de espera: si se rechaza una aprobada, el siguiente en espera es promovido automaticamente con notificacion por email.')

h2('6.6 Fixture / Cuadro de Competencia')
body('El fixture es la estructura competitiva de cada categoria.')
shot('[ CAPTURA DE PANTALLA - Pantalla de fixture ]')
body('Desde el fixture se puede:')
bullet('Visualizar el cuadro completo segun el formato.')
bullet('Registrar resultados de partidos.')
bullet('Editar resultados ya ingresados.')
bullet('Ver el estado de avance de la competencia.')
bullet('Acceder a la agenda con el boton "Ver agenda".')

h3('6.6.1 Generacion del Fixture')
body('Una vez aprobadas las inscripciones, el organizador genera el fixture:')
step('Verificar que las inscripciones esten en estado Aprobada.')
step('Hacer clic en "Generar fixture".')
step('El sistema distribuira las parejas y creara los partidos automaticamente.')
step('Verificar el cuadro generado.')
warn('Una vez generado el fixture no es posible eliminar inscripciones aprobadas ni cambiar el formato de la categoria.')

h3('6.6.2 Registro de Resultados de Partidos')
step('Localizar el partido en el cuadro.')
step('Hacer clic en el icono de lapiz o boton "Ingresar resultado".')
step('Ingresar el marcador set por set (games de cada equipo por set).')
step('Confirmar el resultado.')
body('El sistema actualizara automaticamente el avance en el cuadro.')
note('Se admite registrar walkover (W.O.) o retiro de equipo. El partido se resuelve a favor del equipo presente.')
shot('[ CAPTURA DE PANTALLA - Modal de carga de resultado ]')

h3('6.6.3 Edicion de Resultados')
body(
    'Los resultados ingresados pueden corregirse mediante el icono de lapiz junto a cada partido. '
    'No existe restriccion por estado del partido para realizar correcciones desde el panel del organizador.'
)

h3('6.6.4 Vista por Formato - Eliminacion Simple')
body('Bracket clasico en arbol. Cada ronda avanza al ganador. El cuadro se actualiza automaticamente al registrar resultados.')
shot('[ CAPTURA DE PANTALLA - Bracket eliminacion simple ]')

h3('6.6.5 Vista por Formato - Grupos + Playoff')
body('Fase de grupos con tabla de posiciones (PJ, PG, PP, SG, SP, GG, GP, Pts.) y cuadro de playoff con los clasificados.')
shot('[ CAPTURA DE PANTALLA - Grupos + tabla de posiciones ]')

h3('6.6.6 Vista por Formato - Doble Eliminacion')
body(
    'Cuadro A (ganadores) y Cuadro B (perdedores). El equipo que pierde en A pasa a B. '
    'El perdedor en B queda eliminado. Gran Final entre los ganadores de ambos cuadros.'
)
shot('[ CAPTURA DE PANTALLA - Doble eliminacion ]')

h3('6.6.7 Vista por Formato - Round Robin (Liga)')
body('Todos los equipos se enfrentan entre si. Clasificacion por puntos (3 victoria, 0 derrota). Desempate por diferencia de sets y games.')
shot('[ CAPTURA DE PANTALLA - Tabla Round Robin ]')

h3('6.6.8 Vista por Formato - Americano')
body('Similar al Round Robin pero la clasificacion se determina por games ganados en total. Ideal para jornadas de un dia.')

h3('6.6.9 Vista por Formato - Mexicano')
body(
    'Sistema de rondas dinamicas. Luego de cada ronda, los emparejamientos se recalculan '
    'segun el ranking acumulado. La cantidad de rondas es configurable. '
    'Se muestra tabla de clasificacion general y detalle de cada ronda colapsable.'
)
shot('[ CAPTURA DE PANTALLA - Vista Mexicano ]')

pb()

# ══════════════════════════════════════════════════════════════
#  7. MODULO JUGADORES
# ══════════════════════════════════════════════════════════════
h1('7. MODULO DE JUGADORES (Organizador)')

h2('7.1 Descripcion General')
body(
    'Permite gestionar el directorio de jugadores de la organizacion: '
    'crear nuevos perfiles, consultar y editar informacion existente.'
)
body('Acceso: menu lateral -> Jugadores.')

h2('7.2 Listado de Jugadores')
body('Tabla con todos los jugadores del organizador:')
bullet('Iniciales / Avatar.')
bullet('Nombre y apellido.')
bullet('Correo electronico.')
bullet('DNI.')
bullet('Cantidad de parejas registradas.')
bullet('Acciones: Ver perfil / Editar.')
body('Buscador para filtrar por nombre, apellido, DNI o correo electronico.')
shot('[ CAPTURA DE PANTALLA - Listado de jugadores ]')

h2('7.3 Crear Nuevo Jugador')
body('Hacer clic en "Nuevo jugador" en la esquina superior derecha.')
shot('[ CAPTURA DE PANTALLA - Formulario de nuevo jugador ]')
body('Campos disponibles:')
bullet('Nombre. Obligatorio.')
bullet('Apellido. Obligatorio.')
bullet('Correo electronico. Obligatorio.')
bullet('Telefono de contacto.')
bullet('Fecha de nacimiento.')
bullet('DNI.')
step('Completar los datos del jugador.')
step('Presionar "Guardar".')
body('El jugador quedara registrado y disponible para inscripciones.')

h2('7.4 Editar Jugador')
body(
    'Desde el listado o el perfil del jugador, hacer clic en el icono de lapiz. '
    'Se abre el formulario con los datos pre-cargados para su modificacion.'
)

h2('7.5 Perfil del Jugador')
body('El perfil muestra:')
bullet('Datos personales.')
bullet('Historial de inscripciones y torneos participados.')
bullet('Parejas con las que ha competido.')
shot('[ CAPTURA DE PANTALLA - Perfil del jugador ]')

pb()

# ══════════════════════════════════════════════════════════════
#  8. MODULO INSCRIPCIONES (VISTA GLOBAL)
# ══════════════════════════════════════════════════════════════
h1('8. MODULO DE INSCRIPCIONES (Vista Global)')

h2('8.1 Descripcion General')
body(
    'Ofrece una vista centralizada de todas las solicitudes de inscripcion pendientes '
    'de todos los torneos del organizador, sin necesidad de entrar torneo por torneo.'
)
body('Acceso: menu lateral -> Inscripciones.')

h2('8.2 Panel de Inscripciones Pendientes')
body('Las inscripciones se agrupan por torneo y categoria. Por cada solicitud se muestra:')
bullet('Torneo y categoria.')
bullet('Pareja solicitante (jugador 1 y jugador 2).')
bullet('Fecha de solicitud.')
bullet('Acciones: Aprobar / Rechazar.')
shot('[ CAPTURA DE PANTALLA - Panel de inscripciones pendientes ]')
body('Filtro disponible: por rango de fechas de solicitud.')

h2('8.3 Aprobar o Rechazar Inscripciones')
step('Localizar la inscripcion pendiente en el listado.')
step('Hacer clic en "Aprobar" o "Rechazar".')
step('Confirmar la accion en el dialogo de confirmacion.')
body('Ambas acciones envian notificacion por correo electronico a los jugadores de la pareja.')

h2('8.4 Lista de Espera')
body(
    'Cuando el cupo de una categoria esta completo, las nuevas inscripciones ingresan automaticamente '
    'a la lista de espera. Si se rechaza o cancela una inscripcion aprobada, el sistema promueve '
    'automaticamente a la siguiente pareja de la lista con notificacion por email.'
)
note('La lista de espera es gestionada automaticamente. El organizador puede visualizarla desde las inscripciones de cada categoria.')

pb()

# ══════════════════════════════════════════════════════════════
#  9. MODULO CALENDARIO
# ══════════════════════════════════════════════════════════════
h1('9. MODULO DE CALENDARIO Y AGENDA')

h2('9.1 Descripcion General')
body(
    'Permite programar los partidos en fechas, horarios y canchas especificas. '
    'Ofrece una vista mensual del calendario y herramientas para organizar la agenda de competencia.'
)
body('Acceso: menu lateral -> Calendario.')

h2('9.2 Vista del Calendario')
body('Calendario mensual con los partidos ya programados. Cada partido indica:')
bullet('Equipos participantes.')
bullet('Horario asignado.')
bullet('Sede y cancha.')
body('Navegacion: avanzar y retroceder entre meses. Ver detalle al hacer clic sobre un partido.')
shot('[ CAPTURA DE PANTALLA - Vista del calendario ]')

h2('9.3 Programar Partidos')
body(
    'Los partidos generados por el fixture aparecen sin asignar hasta que el organizador los programe. '
    'Para asignar horario y cancha:'
)
step('Acceder al Calendario o usar el boton "Ver agenda" desde el fixture.')
step('Localizar el partido sin programar en el listado de partidos pendientes.')
step('Seleccionar el dia en el calendario.')
step('Asignar horario de inicio y cancha.')
step('Guardar la programacion.')
shot('[ CAPTURA DE PANTALLA - Programacion de partido ]')

h2('9.4 Asignacion de Canchas')
body(
    'Al programar un partido, se puede seleccionar la cancha de las sedes configuradas. '
    'La cancha puede modificarse en cualquier momento.'
)
note('Las canchas disponibles son las configuradas en Modulo de Configuracion -> Sedes y Canchas.')

pb()

# ══════════════════════════════════════════════════════════════
#  10. MODULO RANKING
# ══════════════════════════════════════════════════════════════
h1('10. MODULO DE RANKING')

h2('10.1 Descripcion General')
body(
    'Permite crear y gestionar tablas de ranking por categoria, '
    'configurar las reglas de distribucion de puntos y visualizar la clasificacion actual.'
)
body('Acceso: menu lateral -> Ranking.')

h2('10.2 Tablas de Ranking')
body('Cada tabla se vincula a una o varias categorias. Se muestra:')
bullet('Posicion de cada pareja/jugador.')
bullet('Puntos acumulados.')
bullet('Cantidad de torneos jugados.')
shot('[ CAPTURA DE PANTALLA - Tabla de ranking ]')

h2('10.3 Crear una Tabla de Ranking')
step('Hacer clic en "Nueva tabla de ranking".')
step('Asignar un nombre a la tabla.')
step('Seleccionar las categorias que la integran.')
step('Configurar las reglas de puntuacion.')
step('Guardar.')

h2('10.4 Reglas de Puntuacion')
body('Para cada tabla se configuran los puntos a otorgar segun el resultado en cada torneo:')
bullet('1er lugar: puntos al ganador del torneo.')
bullet('2do lugar: puntos al finalista.')
bullet('3er lugar: puntos al semifinalista.')
bullet('4to lugar: puntos al semifinalista eliminado.')
bullet('Participacion: puntos por haber participado en el torneo.')
shot('[ CAPTURA DE PANTALLA - Configuracion de reglas de ranking ]')

h2('10.5 Recalculo del Ranking')
body(
    'El boton "Recalcular" actualiza la tabla procesando todos los torneos finalizados '
    'y distribuyendo los puntos segun las reglas configuradas. '
    'Se recomienda ejecutarlo luego de finalizar un torneo o de modificar las reglas.'
)

pb()

# ══════════════════════════════════════════════════════════════
#  11. MODULO REPORTES
# ══════════════════════════════════════════════════════════════
h1('11. MODULO DE REPORTES')

h2('11.1 Descripcion General')
body(
    'Permite obtener informacion consolidada sobre inscripciones, partidos y campeones de torneos. '
    'Los reportes pueden exportarse en formato CSV.'
)
body('Acceso: menu lateral -> Reportes.')

h2('11.2 Reporte de Inscripciones')
body('Listado de equipos inscriptos por torneo y categoria:')
bullet('Torneo y categoria.')
bullet('Pareja (jugador 1 y jugador 2).')
bullet('Estado de la inscripcion.')
bullet('Fecha de inscripcion.')
shot('[ CAPTURA DE PANTALLA - Reporte de inscripciones ]')

h2('11.3 Reporte de Partidos')
body('Listado de partidos con su estado:')
bullet('Torneo y categoria.')
bullet('Equipos participantes.')
bullet('Resultado (si esta completado).')
bullet('Fecha y hora programada.')
bullet('Sede y cancha.')
shot('[ CAPTURA DE PANTALLA - Reporte de partidos ]')

h2('11.4 Reporte de Campeones')
body('Equipos ganadores de cada torneo finalizado por categoria:')
bullet('Torneo y categoria.')
bullet('Pareja campeona (1er lugar).')
bullet('Fecha del torneo.')
shot('[ CAPTURA DE PANTALLA - Reporte de campeones ]')

h2('11.5 Exportacion CSV')
body('Cada seccion de reportes cuenta con el boton "Exportar CSV":')
step('Seleccionar la seccion de reporte deseada.')
step('Hacer clic en "Exportar CSV".')
step('El archivo se descargara automaticamente al dispositivo.')
note('El archivo CSV es compatible con Microsoft Excel, Google Sheets y otras herramientas de analisis.')

pb()

# ══════════════════════════════════════════════════════════════
#  12. MODULO CONFIGURACION
# ══════════════════════════════════════════════════════════════
h1('12. MODULO DE CONFIGURACION')

h2('12.1 Descripcion General')
body(
    'Centraliza todos los parametros de la organizacion. '
    'Se accede a traves de un menu de pestanas con cinco secciones principales.'
)
body('Acceso: menu lateral -> Configuracion.')
shot('[ CAPTURA DE PANTALLA - Panel de configuracion ]')

h2('12.2 Pestana: Datos de la Organizacion')
body('Actualiza la informacion publica del organizador:')
bullet('Nombre de la organizacion.')
bullet('Descripcion.')
bullet('Slug (identificador en URL).')
bullet('Sitio web.')
bullet('Informacion de contacto.')
step('Modificar los campos deseados.')
step('Presionar "Guardar cambios".')

h2('12.3 Pestana: Parametros')
body('Configura los valores predeterminados para nuevos torneos:')
bullet('Maximo de parejas por defecto.')
bullet('Sets por partido por defecto.')
bullet('Games por set por defecto.')
bullet('Precio por pareja por defecto.')

h2('12.4 Pestana: Categorias')
body('Gestiona el catalogo de categorias del organizador (ej.: Primera, Segunda, Mixto, Damas):')
bullet('Crear nueva categoria: ingresar nombre y activarla.')
bullet('Activar / Desactivar categoria: solo las activas estan disponibles al crear torneos.')
shot('[ CAPTURA DE PANTALLA - Gestion de categorias ]')

h2('12.5 Pestana: Colaboradores')
body('Gestiona los usuarios que colaboran con la organizacion con acceso limitado.')
shot('[ CAPTURA DE PANTALLA - Gestion de colaboradores ]')
body('Acciones disponibles:')
bullet('Invitar colaborador: ingresar el email del usuario a invitar.')
bullet('Configurar permisos: definir a que torneos y acciones tiene acceso.')
bullet('Eliminar colaborador: revocar el acceso.')
note('El colaborador debe tener una cuenta existente en el sistema para poder ser agregado.')

h2('12.6 Pestana: Sedes y Canchas')
body('Gestiona las sedes fisicas donde se realizan los torneos y las canchas disponibles.')
shot('[ CAPTURA DE PANTALLA - Gestion de sedes ]')

h3('12.6.1 Crear Sede')
step('Hacer clic en "Nueva sede".')
step('Completar: nombre, direccion, ciudad, telefono, sitio web (opcional).')
step('Guardar.')

h3('12.6.2 Gestionar Canchas de una Sede')
step('Ingresar al detalle de la sede.')
step('Hacer clic en "Agregar cancha".')
step('Ingresar nombre de la cancha y tipo (indoor, outdoor, etc.).')
step('Guardar.')
body('Las canchas pueden editarse o eliminarse desde el listado de la sede.')

pb()

# ══════════════════════════════════════════════════════════════
#  13. MODULO AUDITORIA (ORGANIZADOR)
# ══════════════════════════════════════════════════════════════
h1('13. MODULO DE AUDITORIA (Organizador)')

h2('13.1 Descripcion General')
body(
    'Registra todas las acciones realizadas dentro de la organizacion, '
    'permitiendo trazabilidad completa de quien hizo que y cuando.'
)
body('Acceso: menu lateral -> Auditoria.')

h2('13.2 Filtros Disponibles')
bullet('Por usuario (quien realizo la accion).')
bullet('Por tipo de entidad (torneo, inscripcion, categoria, jugador, etc.).')
bullet('Por rango de fechas.')
shot('[ CAPTURA DE PANTALLA - Log de auditoria ]')

h2('13.3 Informacion Registrada por Evento')
bullet('Fecha y hora de la accion.')
bullet('Usuario responsable.')
bullet('Tipo de accion (crear, modificar, eliminar, aprobar, rechazar, etc.).')
bullet('Entidad afectada.')
bullet('Valores anteriores y nuevos (para modificaciones).')
bullet('Direccion IP del usuario.')

pb()

# ══════════════════════════════════════════════════════════════
#  14. PERFIL DE USUARIO
# ══════════════════════════════════════════════════════════════
h1('14. PERFIL DE USUARIO')

h2('14.1 Descripcion General')
body(
    'Permite a cualquier usuario autenticado consultar y actualizar su informacion personal y de cuenta. '
    'Acceso: hacer clic en el avatar / nombre en la esquina inferior izquierda del menu lateral.'
)
shot('[ CAPTURA DE PANTALLA - Perfil de usuario ]')

h2('14.2 Datos de la Cuenta')
bullet('Correo electronico (no modificable).')
bullet('Fecha de creacion de la cuenta.')
bullet('Rol asignado en el sistema.')

h2('14.3 Perfil de Jugador')
body('Si el usuario tiene un perfil de jugador vinculado, puede consultar y editar:')
bullet('Nombre y apellido.')
bullet('Telefono de contacto.')
bullet('Fecha de nacimiento.')
bullet('DNI.')

h2('14.4 Cambio de Contrasena')
step('Hacer clic en "Cambiar contrasena".')
step('Ingresar la contrasena actual.')
step('Ingresar la nueva contrasena.')
step('Confirmar la nueva contrasena.')
step('Guardar los cambios.')
note('Por seguridad, siempre se debe ingresar la contrasena actual antes de cambiarla.')

pb()

# ══════════════════════════════════════════════════════════════
#  15. PANEL SUPER ADMINISTRADOR
# ══════════════════════════════════════════════════════════════
h1('15. PANEL SUPER ADMINISTRADOR')

h2('15.1 Descripcion General')
body(
    'El Panel Super Administrador es exclusivo para usuarios con rol Super Admin. '
    'Proporciona visibilidad global sobre todo el sistema: organizadores, usuarios, torneos y actividad.'
)
body('Acceso: el enlace "Panel Admin" aparece destacado en rojo en el menu lateral del dashboard.')
shot('[ CAPTURA DE PANTALLA - Panel Super Admin, resumen global ]')

h2('15.2 Resumen Global (Dashboard Admin)')
body('Pantalla principal con metricas del sistema:')
bullet('Organizadores activos.')
bullet('Total de usuarios registrados.')
bullet('Total de torneos creados.')
bullet('Eventos registrados en los ultimos 7 dias.')
bullet('Actividad reciente: log resumido con las ultimas acciones del sistema.')

h2('15.3 Gestion de Organizadores')
body('Administra todos los organizadores registrados en la plataforma:')
bullet('Ver listado completo: nombre, estado, usuarios vinculados, torneos.')
bullet('Activar o desactivar un organizador (los inactivos no pueden operar).')
bullet('Ver los detalles de cada organizacion.')
shot('[ CAPTURA DE PANTALLA - Gestion de organizadores ]')

h2('15.4 Gestion de Usuarios')
body('Administra todos los usuarios del sistema:')
bullet('Ver listado: nombre, email, rol, organizaciones vinculadas, estado, fecha de creacion.')
bullet('Activar o desactivar usuarios.')
bullet('Promover usuario a Super Admin.')
bullet('Degradar Super Admin a usuario estandar.')
shot('[ CAPTURA DE PANTALLA - Gestion de usuarios ]')
warn('La promocion a Super Admin otorga acceso total al sistema. Debe utilizarse con extrema precaucion.')

h2('15.5 Auditoria Global')
body(
    'Vision completa del log de auditoria de toda la plataforma, sin limitacion por organizador. '
    'Permite investigar incidentes y monitorear el comportamiento del sistema.'
)
body('Filtros disponibles:')
bullet('Por usuario.')
bullet('Por organizador.')
bullet('Por tipo de accion.')
bullet('Por rango de fechas.')
shot('[ CAPTURA DE PANTALLA - Auditoria global ]')

pb()

# ══════════════════════════════════════════════════════════════
#  16. BARRA SUPERIOR Y NAVEGACION
# ══════════════════════════════════════════════════════════════
h1('16. BARRA SUPERIOR Y NAVEGACION')

h2('16.1 Descripcion General')
body('La barra superior aparece en todas las pantallas del panel privado y contiene accesos rapidos a funciones comunes.')
shot('[ CAPTURA DE PANTALLA - Barra superior ]')

h2('16.2 Elementos de la Barra Superior')
bullet('Titulo / Nombre de la seccion actual.')
bullet('Indicador de inscripciones pendientes: contador con las inscripciones que requieren atencion del organizador.')
bullet('Menu de usuario: acceso al perfil y cierre de sesion.')

h2('16.3 Menu Lateral (Sidebar)')
body('El menu varia segun el rol del usuario.')

h3('16.3.1 Menu para Organizadores y Colaboradores')
bullet('Dashboard (Inicio).')
bullet('Torneos.')
bullet('Jugadores.')
bullet('Inscripciones.')
bullet('Calendario.')
bullet('Ranking.')
bullet('Reportes.')
bullet('Configuracion.')
bullet('Auditoria.')
bullet('Panel Admin (solo Super Admin, destacado en rojo).')
bullet('Nombre y avatar del usuario en la parte inferior (enlace al perfil).')

h3('16.3.2 Menu para Jugadores')
bullet('Mi Panel.')
bullet('Mi Perfil.')
bullet('Portal Publico.')

pb()

# ══════════════════════════════════════════════════════════════
#  17. SALIR DEL SISTEMA
# ══════════════════════════════════════════════════════════════
h1('17. SALIR DEL SISTEMA')
body('Para cerrar la sesion de forma segura:')
step('Hacer clic en el avatar o nombre de usuario en la esquina inferior del menu lateral.')
step('Seleccionar la opcion "Cerrar sesion".')
step('El sistema cerrara la sesion y redirigira a la pantalla de inicio de sesion.')
warn(
    'Se recomienda cerrar sesion siempre al terminar de trabajar, especialmente en dispositivos compartidos.'
)
shot('[ CAPTURA DE PANTALLA - Opcion de cierre de sesion ]')

pb()

# ══════════════════════════════════════════════════════════════
#  APENDICE A – GLOSARIO
# ══════════════════════════════════════════════════════════════
h1('APENDICE A - GLOSARIO DE TERMINOS')

terms = [
    ('Bracket',        'Cuadro de eliminacion directa que representa el avance de los equipos.'),
    ('Categoria',      'Division de competencia dentro de un torneo (ej. 4ta, 5ta, Mixto, Damas).'),
    ('Colaborador',    'Usuario con acceso limitado a la gestion de una organizacion.'),
    ('Cupo',           'Cantidad maxima de parejas permitidas en una categoria de torneo.'),
    ('DNI',            'Documento Nacional de Identidad. Identificador unico del jugador.'),
    ('Dupla / Pareja', 'Par de jugadores que compiten juntos como equipo en padel.'),
    ('Fixture',        'Estructura competitiva de una categoria. Define partidos y orden de competencia.'),
    ('Games',          'Puntos individuales dentro de un set. Cantidad configurable por el organizador.'),
    ('Gran Final',     'Partido final entre ganadores de Cuadro A y Cuadro B en Doble Eliminacion.'),
    ('Lista de espera','Cola automatica de inscripciones cuando el cupo de una categoria esta completo.'),
    ('Organizador',    'Usuario responsable de crear y gestionar torneos en la plataforma.'),
    ('Playoff',        'Fase eliminatoria posterior a la fase de grupos.'),
    ('Portal Publico', 'Area del sistema accesible sin login para espectadores.'),
    ('Round Robin',    'Formato donde todos los equipos se enfrentan entre si.'),
    ('Sede',           'Lugar fisico donde se realiza el torneo. Contiene una o mas canchas.'),
    ('Set',            'Unidad de juego en padel. Los partidos se definen por un numero de sets configurable.'),
    ('Slug',           'Identificador unico textual de una organizacion, usado en URLs del portal.'),
    ('Super Admin',    'Rol con acceso total al sistema. Gestiona organizadores y usuarios globalmente.'),
    ('Torneo',         'Evento competitivo que agrupa una o varias categorias de padel.'),
    ('Walkover (W.O.)','Resultado asignado cuando un equipo no se presenta. El otro gana por default.'),
]

tbl = doc.add_table(rows=len(terms)+1, cols=2)
tbl.style = 'Table Grid'
table_hdr(tbl, ['Termino', 'Definicion'], fill='1E3A5F')
for i,(t_,d) in enumerate(terms):
    tbl.rows[i+1].cells[0].text = t_
    if tbl.rows[i+1].cells[0].paragraphs[0].runs:
        tbl.rows[i+1].cells[0].paragraphs[0].runs[0].bold = True
    tbl.rows[i+1].cells[1].text = d

doc.add_paragraph()
pb()

# ══════════════════════════════════════════════════════════════
#  APENDICE B – CICLO DE VIDA DEL TORNEO
# ══════════════════════════════════════════════════════════════
h1('APENDICE B - CICLO DE VIDA DE UN TORNEO')
body('El siguiente diagrama describe el ciclo de vida de un torneo y las transiciones de estado posibles:')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    '\n'
    '  [BORRADOR] -> Publicar -> [PUBLICADO]\n'
    '  [PUBLICADO] -> Abrir inscripciones -> [INSCRIPCIONES ABIERTAS]\n'
    '  [INSCRIPCIONES ABIERTAS] <-> Cerrar/Reabrir <-> [INSCRIPCIONES CERRADAS]\n'
    '  [INSCRIPCIONES CERRADAS] -> Iniciar -> [EN CURSO]\n'
    '  [EN CURSO] -> Completar -> [FINALIZADO]\n'
    '\n'
    '  Desde cualquier estado activo -> Cancelar -> [CANCELADO]\n'
)
r.font.name = 'Courier New'
r.font.size = Pt(10)
r.font.color.rgb = C_DARK

doc.add_paragraph()
body('Notas importantes:')
bullet('Un torneo en Borrador no es visible en el Portal Publico.')
bullet('Las inscripciones solo se reciben cuando el estado es Inscripciones Abiertas.')
bullet('El torneo solo puede iniciarse despues de cerrar inscripciones.')
bullet('Un torneo Cancelado no puede reactivarse.')
pb()

# ══════════════════════════════════════════════════════════════
#  APENDICE C – FLUJO DE INSCRIPCION
# ══════════════════════════════════════════════════════════════
h1('APENDICE C - FLUJO COMPLETO DE UNA INSCRIPCION')
body('El siguiente flujo describe el recorrido de una inscripcion desde la solicitud hasta la competencia:')

flow = [
    ('1. Solicitud',   'El jugador envia una solicitud de inscripcion para su pareja desde el portal o el panel del jugador.'),
    ('2. Recepcion',   'La inscripcion aparece en estado Pendiente en el panel del organizador (modulo Inscripciones y modulo Torneos).'),
    ('3. Evaluacion',  'El organizador revisa la solicitud. Si el cupo esta lleno, pasa automaticamente a lista de espera.'),
    ('4a. Aprobacion', 'El organizador aprueba la inscripcion. La pareja recibe notificacion por email con la confirmacion.'),
    ('4b. Rechazo',    'El organizador rechaza la inscripcion. La pareja recibe notificacion por email. Si hay pareja en lista de espera, es promovida automaticamente.'),
    ('5. Fixture',     'La pareja aprobada es incluida en el fixture cuando el organizador lo genera.'),
    ('6. Competencia', 'La pareja participa en los partidos segun el formato de la categoria.'),
]

tbl2 = doc.add_table(rows=len(flow)+1, cols=2)
tbl2.style = 'Table Grid'
table_hdr(tbl2, ['Etapa', 'Descripcion'], fill='22C55E')
for i,(e,d) in enumerate(flow):
    tbl2.rows[i+1].cells[0].text = e
    if tbl2.rows[i+1].cells[0].paragraphs[0].runs:
        tbl2.rows[i+1].cells[0].paragraphs[0].runs[0].bold = True
    tbl2.rows[i+1].cells[1].text = d

# ══════════════════════════════════════════════════════════════
#  GUARDAR
# ══════════════════════════════════════════════════════════════
out = r'C:\proyectos\torneos-padel\MANUAL DE USUARIO - PadelPro.docx'
doc.save(out)
print(f'Manual generado: {out}')
